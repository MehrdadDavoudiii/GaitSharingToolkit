

from __future__ import annotations

import json
import os
import shutil
import base64
import mimetypes
import zipfile
from datetime import datetime
from pathlib import Path

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QGridLayout, QFormLayout,
    QLabel, QPushButton, QLineEdit, QTextEdit, QPlainTextEdit,
    QTableWidget, QTableWidgetItem, QHeaderView, QAbstractItemView,
    QGroupBox, QFrame, QProgressBar, QFileDialog, QMessageBox,
    QDialog, QDialogButtonBox, QCheckBox, QRadioButton,
    QScrollArea, QSplitter, QComboBox, QListWidget, QListWidgetItem,
    QSizePolicy, QSpacerItem, QApplication, QMenu,
    QButtonGroup, QInputDialog,
)
from PySide6.QtCore import Qt, Signal, QThread, QTimer, Slot, QSize
from PySide6.QtGui import QFont, QColor, QPixmap, QPainter, QBrush, QPen

from GaitSharing_config import (
    PALETTE, USER_INFO, USER_PHOTO_PATH, DATA_DIR,
    APP_DIR, load_settings, save_settings,
)
from GaitSharing_database import Database
from GaitSharing_importer import import_dataset, import_selected_folders, sync_dataset
from GaitSharing_excel import create_export_excel
from GaitSharing_parser import REDACTION_TERMS_DEFAULT

try:
    import fitz
except ImportError:
    fitz = None

try:
    from openai import OpenAI as _OpenAI
except ImportError:
    _OpenAI = None

try:
    from docx import Document as _DocxDocument
except ImportError:
    _DocxDocument = None

#  CONSTANTS

REDACTION_TERMS_FILE = APP_DIR / "redaction_terms.json"

TREE_COLS   = ("No.", "ID", "Name", "Exam Date", "Diagnosis",
               "Cond. Left", "Cond. Right", "Measurements", "Model")
TREE_WIDTHS = (45, 110, 170, 95, 260, 120, 120, 150, 110)

DISPLAY_FIELDS = [
    ("Gait Lab ID",    "ganglabor_id"),  ("No. (DB)",       "id"),
    ("Last Name",      "last_name"),     ("First Name",     "first_name"),
    ("Birth Date",     "birth_date"),    ("Gender",         "gender"),
    ("Exam Date",      "exam_date"),     ("Model",          "model"),
    ("Condition L",    "condition_left"),("Condition R",    "condition_right"),
    ("Measurements",   "measurements"),  ("Diagnosis",      "diagnosis"),
    ("Source Folder",  "folder_name"),   ("Import Date",    "import_date"),
]

SEARCH_FIELDS = [
    ("Name (Last)",    "last_name"),   ("First Name",     "first_name"),
    ("Diagnosis",      "diagnosis"),   ("Model",          "model"),
    ("Conditions",     "condition_left"),("Measurements", "measurements"),
    ("Gait Lab ID",    "ganglabor_id"),("Gender",     "gender"),
]

ACK_TEXT = f"""Gait Sharing — Dataset Manager & Analysis Toolkit
===================================================

{USER_INFO['name']}
{USER_INFO['title']}
{USER_INFO['institution']}
{USER_INFO['location']}
{USER_INFO['email']}

Overview
--------
Gait Sharing integrates tools for:
  • Bulk import of multi-language clinical gait-lab PDF reports
  • Subject database with full-text search and keyword tagging
  • Patient dataset export to structured folder trees + Excel summary
  • Clinical PDF anonymisation (PHI redaction via PyMuPDF)
  • AI-assisted clinical report interpretation (OpenAI GPT)
  • C3D data extraction → per-stride kinematics/kinetics/EMG Excel files

Data is stored locally; no network calls are made unless the user
explicitly invokes the AI Interpreter and provides an API key.
"""

LEGAL_WARNING = """Before proceeding, please confirm:

1. All files have been fully anonymised.
2. You have removed all Protected Health Information (PHI).
3. You understand that data will be sent to a third-party AI (OpenAI).
4. You have the legal and ethical right to upload this data for analysis.

You are solely responsible for ensuring data-privacy compliance.
"""

API_HELP = f"""How to get an OpenAI API Key
=============================
1. Visit:  https://platform.openai.com
2. Sign in or create an account.
3. Add a payment method (Pay-As-You-Go).
4. Go to the "API keys" section.
5. Click "Create new secret key" — copy the sk-... value.
6. Paste it into the API Key field in this dialog.

Keep this key private — do not share it.

Support:  {USER_INFO['email']}
"""

PROMPT_PRESETS: dict[str, dict] = {
    "Clinical Biomechanics": {
        "icon": "🦴",
        "description": "Reasoning-based analysis from normal to deviation, with cause-and-effect chains within the stated diagnosis",
        "prompt": """ROLE:
You are a Senior Clinical Biomechanist providing structured biomechanical reasoning on anonymized gait lab data.

DIAGNOSIS CONTEXT:
The patient is being evaluated in the context of: {diagnosis_context}
Frame every observation, every cause, and every downstream effect through the lens of this diagnosis. Do not speculate about conditions outside it.

STRICT PRIVACY DIRECTIVE:
Refer to the subject strictly as "the patient". Do not infer or generate identifiers.

REFERENCE / HEALTHY DATA HANDLING:
If the supplied files contain healthy, normative, control, or pre-injury reference data, you MUST perform a direct, parameter-by-parameter comparison: cite both the patient value and the reference value, and report the deviation in absolute units and as a percent. State explicitly which file is the patient and which is the reference. If no reference data is provided, say so once at the start and use standard published normative ranges as a qualitative anchor only.

ANALYTICAL POSTURE:
Reason — do not label. For each finding, describe (a) how it departs from normal, (b) the most plausible mechanical or neuromuscular driver given the diagnosis, and (c) how it connects upstream and downstream to other findings in the same chain. Build a coherent narrative of cause and effect, not a checklist. Be descriptive but concise. Do NOT prescribe treatment, surgery, orthoses, injections, or medication — the reader is the clinician who decides on management.

OUTPUT FORMAT (markdown):

### 1. Deviation Inventory — Patient vs Normal
A short list (3–6 entries) of the most meaningful departures from normal. For each: parameter, gait phase, patient value, reference value (from supplied healthy data if present, otherwise normative range), absolute and percent deviation.

### 2. Reasoning Chain — From Normal Toward the Observed Pattern
Walk through the deviations one by one, linking them as a chain. For each:
* The mechanical or neuromuscular driver responsible (joint, muscle, lever-arm, motor control).
* Why that driver is consistent with the stated diagnosis.
* Which other deviations in the inventory it most plausibly *causes downstream* or is *caused by upstream*.
Close with a 2–3 sentence synthesis describing the bigger picture: how the originally normal gait has been progressively pulled toward the observed pattern.

### 3. Joint Loading & Functional Consequences
For each joint under abnormal load: loading pattern, supporting moment/power data, and the functional cost (energy expenditure, stability, fall risk). Describe consequences only — do not prescribe.

### 4. Spatiotemporal & Symmetry Reading
Speed, cadence, stride length, stance/swing — interpreted as *evidence* of the chain in section 2, with quantified asymmetry and reasoning about why one side is more affected within the diagnosis.

### 5. Interpretive Confidence & Data Gaps
Where the chain is well supported by the data versus where it is inferred. List specific measurements that would strengthen or weaken the interpretation. (No treatment plan.)
""",
    },

    "Surgical-Relevant Biomechanics": {
        "icon": "🔪",
        "description": "Identifies biomechanically correctable deviations within the diagnosis — descriptive, not prescriptive",
        "prompt": """ROLE:
You are a Clinical Biomechanist describing the biomechanical findings most relevant to a surgical decision conversation. You do NOT recommend procedures, you do NOT prescribe surgery — you describe what is mechanically driving the gait pattern so the surgical team can reason about it.

DIAGNOSIS CONTEXT:
The patient is being evaluated in the context of: {diagnosis_context}
All reasoning must stay within this diagnosis.

STRICT PRIVACY DIRECTIVE:
Refer to the subject strictly as "the patient".

REFERENCE / HEALTHY DATA HANDLING:
If healthy, normative, or pre-injury reference data is among the supplied files, compare directly: cite both patient and reference values, deviation in absolute and percent terms, and identify which file is which. Otherwise rely on published normative ranges as a qualitative anchor only.

ANALYTICAL POSTURE:
Reason about *mechanical correctability* — which deviations appear structural (bone geometry, lever-arm, fixed contracture) versus dynamic (motor control, weakness, spasticity) — without naming procedures, doses, or implants. Be descriptive but concise. Do NOT recommend treatment.

OUTPUT FORMAT (markdown):

### 1. Mechanically Significant Deviations
Short list of deviations that materially shape the gait pattern. For each: joint, plane, gait phase, patient value, reference value (supplied or normative), severity descriptor (mild / moderate / severe).

### 2. Structural vs Dynamic Reasoning
For each deviation, reason whether the data supports a structural origin (e.g., torsional malalignment, lever-arm dysfunction, fixed contracture) or a dynamic origin (e.g., weakness, spasticity, mistimed activation). Cite the kinematic, kinetic, or EMG evidence behind the call.

### 3. Linked Cause and Effect Within the Diagnosis
Tie the deviations together: which finding plausibly drives which, given the diagnosis. Build the chain — identify the lead deviation if one is dominant, and the secondary compensations that follow from it. End with a short synthesis of the bigger picture.

### 4. Implications for Loading & Function
What the chain means for joint loading, energy cost, and stability. Describe — do not prescribe.

### 5. Interpretive Limits
Where additional data (imaging, exam findings, EMG, repeat trials) would change the interpretation. No treatment recommendations.
""",
    },

    "Pediatric / Developmental": {
        "icon": "👶",
        "description": "Pediatric gait reasoning within the stated diagnosis — classification and cause-and-effect, not therapy plans",
        "prompt": """ROLE:
You are a Pediatric Gait Specialist interpreting anonymized gait lab data for a child.

DIAGNOSIS CONTEXT:
The child is being evaluated in the context of: {diagnosis_context}
Frame every observation through this diagnosis. Do not extrapolate to other conditions.

STRICT PRIVACY DIRECTIVE:
Refer to the subject as "the patient" / "the child".

REFERENCE / HEALTHY DATA HANDLING:
If age-matched healthy or normative data is provided in the files, compare directly with patient values — both numbers and percent deviation. Identify which file is patient and which is reference. Otherwise use published age-normative ranges qualitatively.

ANALYTICAL POSTURE:
Reason about *why* the gait pattern looks the way it does within the diagnosis, and how the deviations are mechanically linked. Be descriptive but concise. Do NOT prescribe orthoses, botulinum toxin, surgery, or therapy. Describe the mechanics; the clinical team decides management.

OUTPUT FORMAT (markdown):

### 1. Gait Pattern Classification
Apply the most relevant published classification consistent with the diagnosis (e.g., Winters/Gage/Hicks for hemiplegia; Rodda & Graham — true equinus, jump, apparent equinus, crouch — for diplegia). Support with specific kinematic values.

### 2. Patient vs Normal — Key Departures
Short list: parameter, gait phase, patient value, reference value (from supplied healthy data if present, otherwise normative), absolute and percent deviation.

### 3. Reasoning Chain — One by One, Then the Bigger Picture
Walk through the major deviations and link them:
* The mechanical or neuromuscular driver (lever-arm, contracture, weakness, selective motor control, spasticity).
* Why it is consistent with the diagnosis.
* What it causes downstream and what causes it upstream.
End with a synthesis paragraph showing how the normal developmental gait has been pulled toward the observed pattern.

### 4. Functional Reading
GMFCS-level reasoning (if applicable to the diagnosis), energy expenditure indicators (cadence vs speed), and stability — described as *consequences* of the chain, not as targets for intervention.

### 5. Growth-Linked Considerations
Findings that may evolve with skeletal growth (lever-arm dysfunction, torsional deformities). Describe the trajectory — do not prescribe timing of intervention.

### 6. Interpretive Limits
What additional data would sharpen the interpretation. No therapy or surgical plan.
""",
    },

    "Quick Summary": {
        "icon": "⚡",
        "description": "5-point reasoning summary within the stated diagnosis — concise but causal",
        "prompt": """ROLE:
You are a Clinical Biomechanist providing a rapid reasoning-based summary of gait lab findings.

DIAGNOSIS CONTEXT:
Evaluation is in the context of: {diagnosis_context}
Stay within this diagnosis.

STRICT PRIVACY DIRECTIVE:
Refer to the subject as "the patient".

REFERENCE / HEALTHY DATA HANDLING:
If healthy or normative reference data is supplied, anchor the summary on direct patient-vs-reference comparisons (both numbers, percent deviation). Otherwise note that and use published norms qualitatively.

ANALYTICAL POSTURE:
Be concise but reason — each point should explain *why*, not just *what*. Do NOT prescribe treatment.

OUTPUT FORMAT (markdown):

### Rapid Reasoning Summary
1. **Pattern**: classify the gait pattern in one phrase, consistent with the diagnosis.
2. **Lead deviation**: name it, cite patient value vs reference, and state — in one sentence — what most plausibly drives it within the diagnosis.
3. **Downstream chain**: in one or two sentences, name the secondary compensations the lead deviation appears to produce.
4. **Asymmetry**: present/absent; if present, which side is worse, by how much, and what this suggests within the diagnosis.
5. **Functional consequence**: one sentence on energy/stability/fall-risk implications. (No treatment recommendation.)
""",
    },

    "Comparison Report": {
        "icon": "📊",
        "description": "Direct comparison: patient-vs-healthy, pre/post, or left/right — quantified deltas with causal reasoning",
        "prompt": """ROLE:
You are a Clinical Biomechanist comparing gait data across conditions or between patient and healthy reference.

DIAGNOSIS CONTEXT:
Evaluation is in the context of: {diagnosis_context}
All comparisons should be interpreted through this diagnosis.

STRICT PRIVACY DIRECTIVE:
Refer to the subject as "the patient".

WHAT IS BEING COMPARED:
First identify what the supplied files represent. The most common cases are:
* Patient vs healthy / normative reference — if any file is labeled control, healthy, normative, or reference, treat it as such and structure the entire report as a direct patient-vs-healthy comparison.
* Pre vs post intervention or two time points.
* Left vs right within a single session.
If there is ambiguity, state your assumption explicitly at the start of section 1.

ANALYTICAL POSTURE:
Quantify every difference, then *reason* about why it differs within the diagnosis. Be descriptive but concise. Do NOT prescribe treatment.

OUTPUT FORMAT (markdown):

### 1. Conditions Compared
State the two (or more) conditions, which file represents which, and the direction of expected effect.

### 2. Spatiotemporal Δ
Table: Parameter | Condition A | Condition B | Δ (absolute) | Δ (%) | Clinical relevance.

### 3. Kinematic Δ by Joint
For hip, knee, ankle in each plane: report the change, and reason about whether it represents improvement, deterioration, or compensation given the diagnosis. Highlight changes exceeding MCID where applicable.

### 4. Kinetic Δ
Moment and power changes at the key joints. Reason about whether the patient has shifted strategy (e.g., from ankle to hip power generation) and what that means within the diagnosis.

### 5. Cause-and-Effect Synthesis
Tie the deltas together. If condition A is the patient and condition B is healthy, explicitly walk the reader from the healthy pattern toward the patient's pattern, naming the mechanical drivers in order. If pre vs post, walk through what changed and what most plausibly caused it.

### 6. Overall Reading
Has the patient pattern moved toward, away from, or stayed parallel to normal? What does that imply about the underlying mechanism within the diagnosis? (No treatment recommendation.)
""",
    },
}

# Default prompt key
DEFAULT_PROMPT_KEY = "Clinical Biomechanics"

# For backward compatibility
CLINICAL_PROMPT = PROMPT_PRESETS[DEFAULT_PROMPT_KEY]["prompt"]

def _load_custom_prompts() -> dict[str, dict]:
    settings = load_settings()
    return settings.get("custom_prompts", {})

def _save_custom_prompt(name: str, prompt_text: str) -> None:
    settings = load_settings()
    customs = settings.get("custom_prompts", {})
    customs[name] = {
        "icon": "📝",
        "description": f"Custom prompt: {name}",
        "prompt": prompt_text,
    }
    save_settings({"custom_prompts": customs})

def _delete_custom_prompt(name: str) -> None:
    settings = load_settings()
    customs = settings.get("custom_prompts", {})
    customs.pop(name, None)
    save_settings({"custom_prompts": customs})

#  REUSABLE WIDGETS

def make_page_header(title: str, subtitle: str) -> QWidget:
    bar = QFrame()
    bar.setFixedHeight(56)
    bar.setStyleSheet(f"""
        QFrame {{
            background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                stop:0 {PALETTE['primary']}, stop:1 {PALETTE['primary_lt']});
            border-radius: 0px;
        }}
        QLabel {{ background: transparent; }}
    """)
    lay = QVBoxLayout(bar)
    lay.setContentsMargins(20, 8, 20, 8)
    lay.setSpacing(2)

    t = QLabel(title)
    t.setStyleSheet("color: white; font-size: 15px; font-weight: 700;")
    lay.addWidget(t)

    s = QLabel(subtitle)
    s.setStyleSheet(f"color: {PALETTE['accent']}; font-size: 11px;")
    lay.addWidget(s)

    return bar

def make_accent_btn(text: str) -> QPushButton:
    btn = QPushButton(text)
    btn.setProperty("cssClass", "accent")
    btn.setCursor(Qt.PointingHandCursor)
    btn.style().unpolish(btn)
    btn.style().polish(btn)
    return btn

def make_danger_btn(text: str) -> QPushButton:
    btn = QPushButton(text)
    btn.setProperty("cssClass", "danger")
    btn.setCursor(Qt.PointingHandCursor)
    btn.style().unpolish(btn)
    btn.style().polish(btn)
    return btn

def make_console_log() -> QTextEdit:
    log = QTextEdit()
    log.setReadOnly(True)
    log.setProperty("cssClass", "console")
    log.style().unpolish(log)
    log.style().polish(log)
    return log

def make_subject_table(parent=None) -> QTableWidget:
    table = QTableWidget(parent)
    table.setColumnCount(len(TREE_COLS))
    table.setHorizontalHeaderLabels(TREE_COLS)
    table.setAlternatingRowColors(True)
    table.setSelectionBehavior(QAbstractItemView.SelectRows)
    table.setSelectionMode(QAbstractItemView.SingleSelection)
    table.verticalHeader().setVisible(False)
    table.setEditTriggers(QAbstractItemView.NoEditTriggers)
    table.setSortingEnabled(True)
    table.horizontalHeader().setStretchLastSection(True)
    for i, w in enumerate(TREE_WIDTHS):
        table.setColumnWidth(i, w)
    table.horizontalHeader().setSectionResizeMode(4, QHeaderView.Stretch)
    return table

def populate_subject_table(table: QTableWidget, rows: list[dict]):
    table.setSortingEnabled(False)
    table.setRowCount(len(rows))
    for i, r in enumerate(rows):
        name = f"{r.get('last_name', '')} {r.get('first_name', '')}".strip()
        vals = [
            str(i + 1),
            r.get("ganglabor_id", "") or "",
            name,
            r.get("exam_date", "") or "",
            (r.get("diagnosis") or "")[:80],
            r.get("condition_left", "") or "",
            r.get("condition_right", "") or "",
            r.get("measurements", "") or "",
            r.get("model", "") or "",
        ]
        for col, val in enumerate(vals):
            item = QTableWidgetItem(val)
            item.setData(Qt.UserRole, r["id"])  # store DB id
            if col == 0:
                item.setTextAlignment(Qt.AlignCenter)
            table.setItem(i, col, item)
    table.setSortingEnabled(True)

#  ABOUT DIALOG

class AboutWindow(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("About — Gait Sharing")
        self.setMinimumSize(560, 480)
        self.resize(580, 500)
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        # Header

        # Author card
        card = QWidget()
        card.setStyleSheet(f"background-color: {PALETTE['surface']};")
        clay = QVBoxLayout(card)
        clay.setContentsMargins(24, 16, 24, 16)

        name = QLabel(USER_INFO["name"])
        name.setStyleSheet(f"font-size: 14px; font-weight: 700; color: {PALETTE['primary']};")
        clay.addWidget(name)

        role = QLabel(USER_INFO["title"])
        role.setStyleSheet(f"color: {PALETTE['text_muted']};")
        clay.addWidget(role)

        inst = QLabel(USER_INFO["institution"])
        clay.addWidget(inst)

        email = QLabel(USER_INFO["email"])
        email.setStyleSheet(f"color: {PALETTE['accent']}; font-weight: 600;")
        clay.addWidget(email)
        layout.addWidget(card)

        # Separator
        sep = QFrame()
        sep.setFixedHeight(1)
        sep.setStyleSheet(f"background-color: {PALETTE['border']};")
        layout.addWidget(sep)

        # Text
        txt = QTextEdit()
        txt.setReadOnly(True)
        txt.setPlainText(ACK_TEXT)
        txt.setStyleSheet("border: none; padding: 12px;")
        layout.addWidget(txt)

        # Close button
        btn_lay = QHBoxLayout()
        btn_lay.addStretch()
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(self.close)
        btn_lay.addWidget(close_btn)
        btn_lay.setContentsMargins(0, 0, 16, 12)
        layout.addLayout(btn_lay)

#  TAB 1 — PATIENTS

class PatientsTab(QWidget):
    def __init__(self, db: Database, on_select_cb=None, parent=None):
        super().__init__(parent)
        self.db = db
        self.on_select_cb = on_select_cb
        self._selected_id: int | None = None
        self._build()
        self._start_auto_refresh()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        layout.addWidget(make_page_header(
            "Patients",
            "All imported clinical records — click to view, double-click to edit"))

        # Toolbar
        toolbar = QWidget()
        toolbar.setStyleSheet(f"background-color: {PALETTE['surface']}; border-bottom: 1px solid {PALETTE['border']};")
        tlay = QHBoxLayout(toolbar)
        tlay.setContentsMargins(16, 8, 16, 8)

        edit_btn = QPushButton("✎  Edit")
        edit_btn.clicked.connect(self._edit)
        tlay.addWidget(edit_btn)

        del_btn = QPushButton("✕  Delete")
        del_btn.clicked.connect(self._delete)
        tlay.addWidget(del_btn)

        tlay.addSpacing(12)
        self._sync_flash = QLabel("")
        self._sync_flash.setStyleSheet(f"color: {PALETTE['success']}; font-weight: 700; font-size: 11px;")
        tlay.addWidget(self._sync_flash)

        tlay.addStretch()

        self.count_lbl = QLabel("0 records")
        self.count_lbl.setProperty("cssClass", "muted")
        self.count_lbl.style().unpolish(self.count_lbl)
        self.count_lbl.style().polish(self.count_lbl)
        tlay.addWidget(self.count_lbl)

        layout.addWidget(toolbar)

        # Content area
        content = QWidget()
        content.setStyleSheet(f"background-color: {PALETTE['bg']};")
        clay = QVBoxLayout(content)
        clay.setContentsMargins(16, 12, 16, 12)

        # Table
        self.table = make_subject_table()
        self.table.setSelectionMode(QAbstractItemView.SingleSelection)
        self.table.itemSelectionChanged.connect(self._on_select)
        self.table.doubleClicked.connect(lambda: self._edit())
        clay.addWidget(self.table, stretch=3)

        # Detail panel
        detail_group = QGroupBox("Record Details")
        dgrid = QGridLayout(detail_group)
        dgrid.setSpacing(6)
        dgrid.setContentsMargins(12, 20, 12, 12)

        self._detail_labels: dict[str, QLabel] = {}
        pairs = [
            ("Gait Lab ID",  "ganglabor_id"), ("No.",          "id"),
            ("Name",         "last_name"),    ("First Name",   "first_name"),
            ("Birth Date",   "birth_date"),   ("Gender",       "gender"),
            ("Exam Date",    "exam_date"),    ("Model",        "model"),
            ("Condition L",  "condition_left"),("Condition R", "condition_right"),
            ("Measurements", "measurements"), ("Diagnosis",    "diagnosis"),
        ]
        for i, (label, field) in enumerate(pairs):
            row, col = divmod(i, 2)
            lbl = QLabel(f"{label}:")
            lbl.setStyleSheet("font-weight: 700; font-size: 12px;")
            dgrid.addWidget(lbl, row, col * 2)

            val = QLabel("—")
            val.setWordWrap(True)
            val.setStyleSheet("font-size: 12px;")
            self._detail_labels[field] = val
            dgrid.addWidget(val, row, col * 2 + 1)

        dgrid.setColumnStretch(1, 1)
        dgrid.setColumnStretch(3, 1)
        clay.addWidget(detail_group, stretch=1)

        layout.addWidget(content)

    def refresh(self):
        rows = self.db.get_all()
        populate_subject_table(self.table, rows)
        self.count_lbl.setText(f"{len(rows)} records")

    def _start_auto_refresh(self, interval_ms: int = 2000):
        self._timer = QTimer(self)
        self._timer.timeout.connect(self._check_for_updates)
        self._timer.start(interval_ms)
        self._last_mtime = 0.0

    def _check_for_updates(self):
        try:
            current_mtime = os.path.getmtime(self.db.path)
            if self._last_mtime == 0:
                self._last_mtime = current_mtime
            if current_mtime != self._last_mtime:
                self._last_mtime = current_mtime
                self.refresh()
                self._sync_flash.setText("↻ updated")
                QTimer.singleShot(2000, lambda: self._sync_flash.setText(""))
        except Exception:
            pass

    def load_subjects(self, subjects: list[dict]):
        populate_subject_table(self.table, subjects)
        self.count_lbl.setText(f"{len(subjects)} records")

    def get_selected(self) -> dict | None:
        return self.db.get_by_id(self._selected_id) if self._selected_id else None

    def _get_selected_id(self) -> int | None:
        rows = self.table.selectionModel().selectedRows()
        if not rows:
            return None
        item = self.table.item(rows[0].row(), 0)
        return item.data(Qt.UserRole) if item else None

    def _on_select(self):
        self._selected_id = self._get_selected_id()
        if not self._selected_id:
            return
        r = self.db.get_by_id(self._selected_id)
        if not r:
            return
        for field, lbl in self._detail_labels.items():
            lbl.setText(str(r.get(field) or "—"))
        if self.on_select_cb:
            self.on_select_cb(r)

    def _edit(self):
        if not self._selected_id:
            QMessageBox.information(self, "Select a record", "Click a patient row first.")
            return
        rec = self.db.get_by_id(self._selected_id)
        if rec:
            dlg = EditDialog(self, self.db, rec)
            if dlg.exec():
                self.refresh()

    def _delete(self):
        if not self._selected_id:
            QMessageBox.information(self, "Select a record", "Click a patient row first.")
            return
        rec = self.db.get_by_id(self._selected_id)
        name = (f"{rec.get('last_name','')} {rec.get('first_name','')}".strip()
                or rec["folder_name"])

        reply = QMessageBox.question(
            self, "Confirm Archive",
            f"Archive record for '{name}'?\n\n"
            "The patient will be removed from your list, but original "
            "source files will remain untouched on your hard drive.",
            QMessageBox.Yes | QMessageBox.No,
        )
        if reply == QMessageBox.Yes:
            fp = Path(rec.get("folder_path", ""))
            if fp.exists() and fp.parent == DATA_DIR:
                shutil.rmtree(fp, ignore_errors=True)
            self.db.archive_subject(self._selected_id)
            self._selected_id = None
            self.refresh()
            if self.on_select_cb:
                self.on_select_cb(None)

#  EDIT DIALOG

class EditDialog(QDialog):
    def __init__(self, parent, db: Database, rec: dict):
        super().__init__(parent)
        self.db = db
        self.rec = rec
        gait_id = rec.get("ganglabor_id") or rec.get("folder_name", "")
        self.setWindowTitle(f"Edit — {gait_id}")
        self.resize(640, 600)
        self.setMinimumWidth(520)
        self._original: dict[str, str] = {}
        self._widgets: dict[str, QWidget] = {}
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        # Header
        hdr = QFrame()
        hdr.setFixedHeight(48)
        hdr.setStyleSheet(f"background-color: {PALETTE['primary']};")
        hlay = QHBoxLayout(hdr)
        hlay.setContentsMargins(20, 0, 20, 0)
        t = QLabel("Edit Patient Record")
        t.setStyleSheet("color: white; font-size: 14px; font-weight: 700;")
        hlay.addWidget(t)
        layout.addWidget(hdr)

        # Scroll area for fields
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("border: none;")

        form_widget = QWidget()
        form_layout = QFormLayout(form_widget)
        form_layout.setContentsMargins(24, 16, 24, 16)
        form_layout.setSpacing(10)

        editable = [
            f for f in DISPLAY_FIELDS
            if f[1] not in ("id", "folder_name", "import_date",
                            "raw_pdf_text", "pdf_path", "folder_path")
        ]

        for label, field in editable:
            val = self.rec.get(field) or ""
            self._original[field] = val

            if field in ("diagnosis", "measurements"):
                w = QTextEdit()
                w.setPlainText(val)
                w.setMaximumHeight(80)
                self._widgets[field] = w
            else:
                w = QLineEdit(val)
                self._widgets[field] = w

            lbl = QLabel(f"{label}:")
            lbl.setStyleSheet("font-weight: 700;")
            form_layout.addRow(lbl, w)

        scroll.setWidget(form_widget)
        layout.addWidget(scroll)

        # Button bar
        btn_bar = QWidget()
        btn_bar.setStyleSheet(f"background-color: {PALETTE['bg']}; border-top: 1px solid {PALETTE['border']}; color: {PALETTE['primary']};")
        blay = QHBoxLayout(btn_bar)
        blay.setContentsMargins(16, 10, 16, 10)

        blay.addStretch()

        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        blay.addWidget(cancel_btn)

        revert_btn = QPushButton("↺  Revert")
        revert_btn.clicked.connect(self._revert)
        blay.addWidget(revert_btn)

        save_btn = make_accent_btn("Save")
        save_btn.clicked.connect(self._save)
        blay.addWidget(save_btn)

        layout.addWidget(btn_bar)

    def _revert(self):
        for field, w in self._widgets.items():
            val = self._original[field]
            if isinstance(w, QTextEdit):
                w.setPlainText(val)
            else:
                w.setText(val)

    def _save(self):
        data = {}
        for field, w in self._widgets.items():
            if isinstance(w, QTextEdit):
                data[field] = w.toPlainText().strip()
            else:
                data[field] = w.text().strip()
        self.db.update_subject(self.rec["id"], data)
        self.accept()

#  TAB 2 — IMPORT

class ImportWorker(QThread):
    progress = Signal(int, int, str)
    finished = Signal(int, int, int, list)

    def __init__(self, folder, db, direct_mode=True):
        super().__init__()
        self.folder = folder
        self.db = db
        self.direct_mode = direct_mode

    def run(self):
        import_dataset(
            self.folder, self.db,
            progress_cb=self._on_progress,
            done_cb=self._on_done,
            direct_mode=self.direct_mode,
        )

    def _on_progress(self, current, total, name):
        self.progress.emit(current, total, name)

    def _on_done(self, imported, updated, skipped, failed):
        self.finished.emit(imported, updated, skipped, failed)

class SelectedFolderWorker(QThread):
    progress = Signal(int, int, str)
    finished = Signal(int, int, int, list)

    def __init__(self, folders, db, direct_mode=True):
        super().__init__()
        self.folders = folders
        self.db = db
        self.direct_mode = direct_mode

    def run(self):
        import_selected_folders(
            self.folders, self.db,
            progress_cb=self._on_progress,
            done_cb=self._on_done,
            direct_mode=self.direct_mode,
        )

    def _on_progress(self, current, total, name):
        self.progress.emit(current, total, name)

    def _on_done(self, imported, updated, skipped, failed):
        self.finished.emit(imported, updated, skipped, failed)

class ImportTab(QWidget):
    def __init__(self, db: Database, refresh_patients_cb, parent=None):
        super().__init__(parent)
        self.db = db
        self.refresh_patients_cb = refresh_patients_cb
        self._worker = None

        settings = load_settings()
        saved = settings.get("source_folder", "")
        self._dataset_folder: Path | None = Path(saved) if saved else None
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        layout.addWidget(make_page_header(
            "Import & Sync",
            "Connect to your dataset and synchronize changes"))

        # Content
        content = QWidget()
        clay = QVBoxLayout(content)
        clay.setContentsMargins(16, 12, 16, 12)
        clay.setSpacing(10)

        # Source folder group
        src_group = QGroupBox("Dataset Root Folder (source of truth)")
        sglay = QHBoxLayout(src_group)
        self.folder_lbl = QLineEdit(str(self._dataset_folder or ""))
        self.folder_lbl.setReadOnly(True)
        self.folder_lbl.setPlaceholderText("Select a dataset root folder...")
        sglay.addWidget(self.folder_lbl)
        browse_btn = QPushButton("Browse…")
        browse_btn.clicked.connect(self._browse)
        sglay.addWidget(browse_btn)
        clay.addWidget(src_group)

        hint = QLabel("The folder should contain one subfolder per subject. "
                       "Full Sync detects both new and deleted subjects.")
        hint.setProperty("cssClass", "muted")
        hint.style().unpolish(hint)
        hint.style().polish(hint)
        hint.setWordWrap(True)
        clay.addWidget(hint)

        # Action buttons
        btn_row = QHBoxLayout()
        self.import_btn = make_accent_btn("▶  Import All")
        self.import_btn.clicked.connect(self._start)
        btn_row.addWidget(self.import_btn)

        self.sync_btn = QPushButton("🔄  Full Sync")
        self.sync_btn.clicked.connect(self._full_sync)
        btn_row.addWidget(self.sync_btn)

        folder_btn = QPushButton("📁  Import / Update Folders")
        folder_btn.clicked.connect(self._import_single_folder)
        btn_row.addWidget(folder_btn)

        manual_btn = QPushButton("➕  Add Manually")
        manual_btn.clicked.connect(self._open_manual_add)
        btn_row.addWidget(manual_btn)

        btn_row.addStretch()
        clay.addLayout(btn_row)

        # Progress
        prog_group = QGroupBox("Progress")
        pglay = QVBoxLayout(prog_group)
        self.status_lbl = QLabel("Ready.")
        pglay.addWidget(self.status_lbl)
        self.pbar = QProgressBar()
        self.pbar.setMaximum(100)
        pglay.addWidget(self.pbar)
        clay.addWidget(prog_group)

        # Log
        log_group = QGroupBox("Log")
        lglay = QVBoxLayout(log_group)
        self.log = make_console_log()
        self.log.setMinimumHeight(100)
        lglay.addWidget(self.log)
        clay.addWidget(log_group, stretch=1)

        # Wrap in scroll area for small screens (14")
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setWidget(content)
        scroll.setStyleSheet("QScrollArea { border: none; }")
        layout.addWidget(scroll)

    def _log(self, msg: str):
        self.log.append(msg)

    def _browse(self):
        path = QFileDialog.getExistingDirectory(self, "Select Dataset Root Folder")
        if path:
            self._dataset_folder = Path(path)
            self.folder_lbl.setText(str(self._dataset_folder))
            save_settings({"source_folder": str(self._dataset_folder)})
            n = sum(1 for p in self._dataset_folder.iterdir() if p.is_dir())
            self._log(f"Source folder set: {path}  ({n} subfolders)")

    def _check_ready(self) -> bool:
        if not self._dataset_folder or not self._dataset_folder.exists():
            QMessageBox.critical(self, "No Folder",
                                 "Please select a valid Dataset Root Folder first.")
            return False
        if fitz is None:
            QMessageBox.critical(self, "Missing Library",
                                 "PyMuPDF (fitz) required — pip install pymupdf")
            return False
        return True

    def _set_buttons_enabled(self, enabled: bool):
        self.import_btn.setEnabled(enabled)
        self.sync_btn.setEnabled(enabled)

    def _start(self):
        if not self._check_ready():
            return
        self._set_buttons_enabled(False)
        self._log(f"\n{'─'*55}\nImport All  [Direct Mode — no copying]\n"
                  f"Source: {self._dataset_folder}\n{'─'*55}")

        self._worker = ImportWorker(self._dataset_folder, self.db)
        self._worker.progress.connect(self._on_progress)
        self._worker.finished.connect(self._on_done)
        self._worker.start()

    def _on_progress(self, current, total, folder):
        if total == 0:
            return
        self.pbar.setValue(int((current / total) * 100))
        self.status_lbl.setText(f"Processing {current}/{total}: {folder}")
        self._log(f"  [{current}/{total}]  {folder}")

    def _on_done(self, imported, updated, skipped, failed):
        self.pbar.setValue(100)
        self.status_lbl.setText(
            f"Done — {imported} new, {updated} updated, "
            f"{skipped} skipped (no PDF), {len(failed)} errors")
        self._log(
            f"\n{'─'*55}\n"
            f"COMPLETE  {imported} new  {updated} updated  "
            f"{skipped} skipped  {len(failed)} errors")
        if failed:
            self._log("Errors:")
            for name, reason in failed:
                self._log(f"  ✗ {name}: {reason}")
        self._log(f"{'─'*55}\n")
        self._set_buttons_enabled(True)
        self.refresh_patients_cb()

    def _full_sync(self):
        if not self._check_ready():
            return
        self._log(f"\n{'─'*55}\nFull Sync scan…\n{'─'*55}")
        try:
            new_folders, deleted_names = sync_dataset(self._dataset_folder, self.db)
        except Exception as exc:
            QMessageBox.critical(self, "Sync Error", str(exc))
            return

        if not new_folders and not deleted_names:
            self._log("Sync: database is already up-to-date with source folder.")
            QMessageBox.information(
                self, "Already Up-to-Date",
                f"No changes detected.\n\nAll subject folders are present in:\n"
                f"{self._dataset_folder}")
            return

        self._log(f"Sync found: {len(new_folders)} new, {len(deleted_names)} missing.")

        dlg = FullSyncDialog(self, new_folders, deleted_names)
        if dlg.exec():
            to_import = dlg.selected_new
            to_delete = dlg.selected_delete

            if to_delete:
                for name in to_delete:
                    try:
                        rec = self.db.get_by_folder(name)
                        if rec:
                            fp = Path(rec.get("folder_path", ""))
                            if fp.exists() and fp.parent == DATA_DIR:
                                shutil.rmtree(fp, ignore_errors=True)
                            self.db.archive_subject(rec["id"])
                            self._log(f"  📦  Archived: {name}")
                    except Exception as e:
                        self._log(f"  ✗  Error archiving {name}: {e}")
                self.refresh_patients_cb()

            if to_import:
                self._set_buttons_enabled(False)
                self._log(f"Importing {len(to_import)} new folder(s)  [Direct Mode]…")
                self._worker = SelectedFolderWorker(to_import, self.db)
                self._worker.progress.connect(self._on_progress)
                self._worker.finished.connect(self._on_done)
                self._worker.start()

    def _import_single_folder(self):
        if not self._check_ready():
            return
        parent_path = QFileDialog.getExistingDirectory(
            self, "Select folder containing subject folder(s)")
        if not parent_path:
            return
        parent = Path(parent_path)
        subfolders = sorted(
            [p for p in parent.iterdir() if p.is_dir()],
            key=lambda p: p.name,
        )
        if not subfolders:
            subfolders = [parent]

        existing_folders = {r["folder_name"] for r in self.db.get_all()}
        dlg = UpdateDialog(self, subfolders, existing_folders)
        if dlg.exec():
            if dlg.selected:
                self._set_buttons_enabled(False)
                self._log(f"\n{'─'*55}\nImporting {len(dlg.selected)} folder(s)  [Direct Mode]…\n{'─'*55}")
                self._worker = SelectedFolderWorker(dlg.selected, self.db)
                self._worker.progress.connect(self._on_progress)
                self._worker.finished.connect(self._on_done)
                self._worker.start()

    def _open_manual_add(self):
        dlg = ManualSubjectDialog(self)
        if dlg.exec():
            if dlg.result:
                try:
                    existing = self.db.get_by_folder(dlg.result["folder_name"])
                    self.db.upsert_subject(dlg.result)
                    action = "updated" if existing else "added"
                    self._log(f"✔ Subject '{dlg.result['folder_name']}' {action}.")
                    self.refresh_patients_cb()
                    QMessageBox.information(
                        self, "Saved",
                        f"Subject '{dlg.result['folder_name']}' {action} successfully.")
                except Exception as exc:
                    QMessageBox.critical(self, "Error", str(exc))

#  SYNC / UPDATE / MANUAL DIALOGS

class FullSyncDialog(QDialog):
    def __init__(self, parent, new_folders: list[Path], deleted_names: list[str]):
        super().__init__(parent)
        self.setWindowTitle(f"Full Sync — {len(new_folders)} new, {len(deleted_names)} missing")
        self.resize(700, 560)
        self.setMinimumSize(540, 380)
        self.selected_new: list[Path] = []
        self.selected_delete: list[str] = []
        self._new = new_folders
        self._deleted = deleted_names
        self._new_cbs: list[QCheckBox] = []
        self._del_cbs: list[QCheckBox] = []
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        # Header
        hdr = QFrame()
        hdr.setFixedHeight(50)
        hdr.setStyleSheet(f"background-color: {PALETTE['primary']};")
        hlay = QHBoxLayout(hdr)
        hlay.setContentsMargins(20, 0, 20, 0)
        t = QLabel(f"🔄  Full Sync — ✚ {len(self._new)} new   📦 {len(self._deleted)} missing")
        t.setStyleSheet("color: white; font-size: 13px; font-weight: 700;")
        hlay.addWidget(t)
        layout.addWidget(hdr)

        # Controls
        ctrl = QHBoxLayout()
        ctrl.setContentsMargins(16, 8, 16, 4)
        self._count_lbl = QLabel()
        self._count_lbl.setStyleSheet(f"font-weight: 700; color: {PALETTE['primary']};")
        ctrl.addWidget(self._count_lbl)
        ctrl.addStretch()
        sel_all = QPushButton("Select All")
        sel_all.clicked.connect(self._select_all)
        ctrl.addWidget(sel_all)
        desel = QPushButton("Deselect All")
        desel.clicked.connect(self._deselect_all)
        ctrl.addWidget(desel)
        layout.addLayout(ctrl)

        # Scrollable content
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("border: none;")
        body = QWidget()
        blay = QVBoxLayout(body)
        blay.setContentsMargins(16, 0, 16, 0)
        blay.setSpacing(0)

        if self._new:
            sec = QLabel(f"  ✚  {len(self._new)} NEW — present in source, not in database")
            sec.setFixedHeight(28)
            sec.setStyleSheet(f"background-color: {PALETTE['success']}; color: white; font-weight: 700; font-size: 11px; border-radius: 4px;")
            blay.addWidget(sec)
            for folder in self._new:
                cb = QCheckBox(folder.name)
                cb.setChecked(True)
                cb.toggled.connect(self._update_count)
                self._new_cbs.append(cb)
                blay.addWidget(cb)

        if self._deleted:
            blay.addSpacing(8)
            sec = QLabel(f"  📦  {len(self._deleted)} MISSING — ready to be archived")
            sec.setFixedHeight(28)
            sec.setStyleSheet(f"background-color: {PALETTE['warning']}; color: white; font-weight: 700; font-size: 11px; border-radius: 4px;")
            blay.addWidget(sec)
            for name in self._deleted:
                cb = QCheckBox(name)
                cb.setChecked(True)
                cb.toggled.connect(self._update_count)
                self._del_cbs.append(cb)
                blay.addWidget(cb)

        blay.addStretch()
        scroll.setWidget(body)
        layout.addWidget(scroll)

        # Buttons
        btn_bar = QWidget()
        btn_bar.setStyleSheet(f"background-color: {PALETTE['accent']}; border-top: 1px solid {PALETTE['accent_dk']}; color: {PALETTE['primary']};")
        bblay = QHBoxLayout(btn_bar)
        bblay.setContentsMargins(16, 10, 16, 10)
        self._apply_btn = make_accent_btn("▶  Apply")
        self._apply_btn.clicked.connect(self._on_apply)
        bblay.addWidget(self._apply_btn)
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        bblay.addWidget(cancel_btn)
        bblay.addStretch()
        layout.addWidget(btn_bar)

        self._update_count()

    def _update_count(self):
        n_imp = sum(cb.isChecked() for cb in self._new_cbs)
        n_del = sum(cb.isChecked() for cb in self._del_cbs)
        self._count_lbl.setText(f"✚ {n_imp} to import   📦 {n_del} to archive")
        self._apply_btn.setEnabled(n_imp + n_del > 0)
        self._apply_btn.setText(f"▶  Apply  (import {n_imp}, archive {n_del})")

    def _select_all(self):
        for cb in self._new_cbs + self._del_cbs:
            cb.setChecked(True)

    def _deselect_all(self):
        for cb in self._new_cbs + self._del_cbs:
            cb.setChecked(False)

    def _on_apply(self):
        self.selected_new = [f for f, cb in zip(self._new, self._new_cbs) if cb.isChecked()]
        self.selected_delete = [n for n, cb in zip(self._deleted, self._del_cbs) if cb.isChecked()]
        self.accept()

class UpdateDialog(QDialog):
    def __init__(self, parent, all_folders: list[Path], existing_names: set[str]):
        super().__init__(parent)
        self.setWindowTitle(f"Select Folders ({len(all_folders)} found)")
        self.resize(680, 520)
        self.selected: list[Path] = []
        self._folders = all_folders
        self._existing = existing_names
        self._cbs: list[QCheckBox] = []
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        n_new = sum(1 for f in self._folders if f.name not in self._existing)
        n_upd = len(self._folders) - n_new

        hdr = QFrame()
        hdr.setFixedHeight(50)
        hdr.setStyleSheet(f"background-color: {PALETTE['primary']};")
        hlay = QHBoxLayout(hdr)
        hlay.setContentsMargins(20, 0, 20, 0)
        t = QLabel(f"📁  {len(self._folders)} folder(s) — ✚ {n_new} new  ⚠ {n_upd} update")
        t.setStyleSheet("color: white; font-size: 13px; font-weight: 700;")
        hlay.addWidget(t)
        layout.addWidget(hdr)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("border: none;")
        body = QWidget()
        blay = QVBoxLayout(body)
        blay.setContentsMargins(16, 8, 16, 8)

        for folder in self._folders:
            is_existing = folder.name in self._existing
            tag = " ⚠ UPDATE" if is_existing else " ✚ NEW"
            cb = QCheckBox(f"{folder.name}{tag}")
            cb.setChecked(not is_existing)
            self._cbs.append(cb)
            blay.addWidget(cb)

        blay.addStretch()
        scroll.setWidget(body)
        layout.addWidget(scroll)

        btn_bar = QWidget()
        btn_bar.setStyleSheet(f"background-color: {PALETTE['accent']}; border-top: 1px solid {PALETTE['accent_dk']}; color: {PALETTE['primary']};")
        bblay = QHBoxLayout(btn_bar)
        bblay.setContentsMargins(16, 10, 16, 10)
        ok_btn = make_accent_btn("▶  Import / Update Selected")
        ok_btn.clicked.connect(self._on_ok)
        bblay.addWidget(ok_btn)
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        bblay.addWidget(cancel_btn)
        bblay.addStretch()
        layout.addWidget(btn_bar)

    def _on_ok(self):
        self.selected = [f for f, cb in zip(self._folders, self._cbs) if cb.isChecked()]
        self.accept()

class ManualSubjectDialog(QDialog):
    _FIELDS = [
        ("ganglabor_id",    "Gait Lab ID *",          "entry"),
        ("last_name",       "Last Name",              "entry"),
        ("first_name",      "First Name",             "entry"),
        ("birth_date",      "Birth Date (DD.MM.YYYY)","entry"),
        ("gender",          "Gender",                 "entry"),
        ("exam_date",       "Exam Date (DD.MM.YYYY)", "entry"),
        ("diagnosis",       "Diagnosis",              "text"),
        ("condition_left",  "Condition Left",         "entry"),
        ("condition_right", "Condition Right",        "entry"),
        ("measurements",    "Measurements",           "entry"),
        ("model",           "Model",                  "entry"),
    ]

    def __init__(self, parent):
        super().__init__(parent)
        self.setWindowTitle("Add Subject Manually")
        self.resize(560, 520)
        self.result: dict | None = None
        self._widgets: dict[str, QWidget] = {}
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        hdr = QFrame()
        hdr.setFixedHeight(46)
        hdr.setStyleSheet(f"background-color: {PALETTE['primary']};")
        hlay = QHBoxLayout(hdr)
        hlay.setContentsMargins(20, 0, 20, 0)
        t = QLabel("➕  Add Subject Manually")
        t.setStyleSheet("color: white; font-size: 13px; font-weight: 700;")
        hlay.addWidget(t)
        layout.addWidget(hdr)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("border: none;")
        form_w = QWidget()
        form = QFormLayout(form_w)
        form.setContentsMargins(24, 16, 24, 16)
        form.setSpacing(8)

        for col, label, wtype in self._FIELDS:
            lbl = QLabel(f"{label}:")
            lbl.setStyleSheet("font-weight: 700;")
            if wtype == "text":
                w = QTextEdit()
                w.setMaximumHeight(70)
            else:
                w = QLineEdit()
            self._widgets[col] = w
            form.addRow(lbl, w)

        scroll.setWidget(form_w)
        layout.addWidget(scroll)

        btn_bar = QWidget()
        btn_bar.setStyleSheet(f"background-color: {PALETTE['bg']}; border-top: 1px solid {PALETTE['border']}; color: {PALETTE['primary']};")
        bblay = QHBoxLayout(btn_bar)
        bblay.setContentsMargins(16, 10, 16, 10)
        save_btn = make_accent_btn("✔  Save Subject")
        save_btn.clicked.connect(self._on_save)
        bblay.addWidget(save_btn)
        clear_btn = QPushButton("↺  Clear Form")
        clear_btn.clicked.connect(self._clear)
        bblay.addWidget(clear_btn)
        cancel_btn = QPushButton("✖  Cancel")
        cancel_btn.clicked.connect(self.reject)
        bblay.addWidget(cancel_btn)
        bblay.addStretch()
        layout.addWidget(btn_bar)

    def _clear(self):
        for col, label, wtype in self._FIELDS:
            w = self._widgets[col]
            if isinstance(w, QTextEdit):
                w.clear()
            else:
                w.clear()

    def _on_save(self):
        data = {}
        for col, label, wtype in self._FIELDS:
            w = self._widgets[col]
            if isinstance(w, QTextEdit):
                data[col] = w.toPlainText().strip()
            else:
                data[col] = w.text().strip()

        if not data.get("ganglabor_id"):
            QMessageBox.warning(self, "Required", "Gait Lab ID is required.")
            return

        data["folder_name"] = data["ganglabor_id"]
        data["folder_path"] = ""
        data["import_date"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        self.result = data
        self.accept()

#  TAG ENTRY WIDGET (for Search)

class TagEntry(QWidget):

    def __init__(self, label: str, field: str, parent=None):
        super().__init__(parent)
        self.field = field
        self.keywords: list[str] = []

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(2)

        row = QHBoxLayout()
        lbl = QLabel(f"{label}:")
        lbl.setFixedWidth(110)
        lbl.setStyleSheet("font-weight: 700; font-size: 12px;")
        row.addWidget(lbl)

        self.entry = QLineEdit()
        self.entry.setPlaceholderText("Type keyword…")
        self.entry.returnPressed.connect(self._add)
        self.entry.setMaximumWidth(250)
        row.addWidget(self.entry)

        #row.addWidget(self.entry)
        

        add_btn = QPushButton("+ Add")
        add_btn.setFixedWidth(90)
        add_btn.clicked.connect(self._add)
        row.addWidget(add_btn)

        layout.addLayout(row)

        self.tag_area = QHBoxLayout()
        self.tag_area.setContentsMargins(115, 0, 0, 0)
        layout.addLayout(self.tag_area)

    def _add(self):
        kw = self.entry.text().strip()
        if kw and kw not in self.keywords:
            self.keywords.append(kw)
            self.entry.clear()
            self._redraw()

    def _redraw(self):
        while self.tag_area.count():
            item = self.tag_area.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        for kw in self.keywords:
            chip = QPushButton(f"  {kw}  ×")
            chip.setFixedHeight(24)
            chip.setCursor(Qt.PointingHandCursor)
            chip.setStyleSheet(f"""
                QPushButton {{
                    background-color: {PALETTE['tag_bg']};
                    color: {PALETTE['tag_fg']};
                    border: 1px solid {PALETTE['border']};
                    border-radius: 12px;
                    font-size: 11px;
                    padding: 0 8px;
                }}
                QPushButton:hover {{
                    background-color: #D4E0A8;
                }}
            """)
            chip.clicked.connect(lambda checked, k=kw: self._remove(k))
            self.tag_area.addWidget(chip)

        self.tag_area.addStretch()

    def _remove(self, kw: str):
        self.keywords = [k for k in self.keywords if k != kw]
        self._redraw()

    def clear(self):
        self.keywords = []
        self.entry.clear()
        self._redraw()

    def get(self) -> list[str]:
        return self.keywords.copy()

#  TAB 3 — SEARCH

class SearchTab(QWidget):
    def __init__(self, db: Database, patients_tab: PatientsTab, parent=None):
        super().__init__(parent)
        self.db = db
        self.patients_tab = patients_tab
        self._results: list[dict] = []
        self._export_tab = None
        self._build()

    def set_export_tab(self, tab):
        self._export_tab = tab

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        layout.addWidget(make_page_header("Search", "Add keywords to filter subjects"))

        content = QWidget()
        clay = QVBoxLayout(content)
        clay.setContentsMargins(16, 12, 16, 12)
        clay.setSpacing(8)

        # Tag entries grid
        grid_w = QWidget()
        grid = QGridLayout(grid_w)
        grid.setContentsMargins(0, 0, 0, 0)
        grid.setSpacing(4)
        self._tag_entries: dict[str, TagEntry] = {}
        for i, (label, field) in enumerate(SEARCH_FIELDS):
            row, col = divmod(i, 2)
            te = TagEntry(label, field)
            grid.addWidget(te, row, col)
            self._tag_entries[field] = te
        clay.addWidget(grid_w)

        # Date ranges
        date_row = QHBoxLayout()
        date_row.setSpacing(6)
        date_lbl = QLabel("Exam Date from:")
        date_lbl.setStyleSheet("font-weight: 700; font-size: 12px;")
        date_lbl.setFixedWidth(100)
        date_row.addWidget(date_lbl)
        self.date_from = QLineEdit()
        self.date_from.setPlaceholderText("YYYY-MM-DD")
        self.date_from.setFixedWidth(120)
        date_row.addWidget(self.date_from)
        to_lbl = QLabel("to:")
        to_lbl.setStyleSheet("font-weight: 700; font-size: 12px;")
        date_row.addWidget(to_lbl)
        self.date_to = QLineEdit()
        self.date_to.setPlaceholderText("YYYY-MM-DD")
        self.date_to.setFixedWidth(120)
        date_row.addWidget(self.date_to)
        date_row.addStretch()
        clay.addLayout(date_row)

        # Logic + actions
        action_row = QHBoxLayout()
        action_row.setSpacing(6)
        logic_lbl = QLabel("Field logic:")
        logic_lbl.setStyleSheet("font-weight: 700; font-size: 12px;")
        logic_lbl.setFixedWidth(100)
        action_row.addWidget(logic_lbl)
        self.logic_and = QRadioButton("AND (all must match)")
        self.logic_and.setChecked(True)
        action_row.addWidget(self.logic_and)
        self.logic_or = QRadioButton("OR (any matches)")
        action_row.addWidget(self.logic_or)
        action_row.addStretch()
        clear_btn = QPushButton("Clear All")
        clear_btn.clicked.connect(self._clear)
        action_row.addWidget(clear_btn)

        search_btn = make_accent_btn("🔍  Search")
        search_btn.clicked.connect(self._search)
        action_row.addWidget(search_btn)
        clay.addLayout(action_row)

        # Results
        results_group = QGroupBox("Results")
        rglay = QVBoxLayout(results_group)

        self.result_count_lbl = QLabel("No search run yet.")
        self.result_count_lbl.setStyleSheet("font-weight: 700;")
        rglay.addWidget(self.result_count_lbl)

        self.res_table = make_subject_table()
        self.res_table.setSelectionMode(QAbstractItemView.ExtendedSelection)
        rglay.addWidget(self.res_table)

        btn_row = QHBoxLayout()
        send_btn = QPushButton("Send search results to Export →")
        send_btn.clicked.connect(self._send_to_export)
        btn_row.addWidget(send_btn)

        send_all_btn = QPushButton("📤  Export entire database →")
        send_all_btn.clicked.connect(self._send_all_to_export)
        btn_row.addWidget(send_all_btn)
        btn_row.addStretch()
        rglay.addLayout(btn_row)

        clay.addWidget(results_group, stretch=1)
        layout.addWidget(content)

    def _search(self):
        filters = {}
        for field, te in self._tag_entries.items():
            kws = te.get()
            if kws:
                filters[field] = kws
        df = self.date_from.text().strip()
        dt = self.date_to.text().strip()
        if df:
            filters["_date_from"] = [df]
        if dt:
            filters["_date_to"] = [dt]

        logic = "AND" if self.logic_and.isChecked() else "OR"
        self._results = self.db.search(filters, logic)
        populate_subject_table(self.res_table, self._results)
        self.result_count_lbl.setText(f"{len(self._results)} patient(s) found.")

    def _clear(self):
        for te in self._tag_entries.values():
            te.clear()
        self.date_from.clear()
        self.date_to.clear()
        self.res_table.setRowCount(0)
        self._results = []
        self.result_count_lbl.setText("No search run yet.")

    def _send_to_export(self):
        rows = self._results if self._results else []
        if self._export_tab and rows:
            self._export_tab.load_subjects(rows)
            QMessageBox.information(self, "Sent",
                                    f"{len(rows)} patient(s) sent to Export tab.")

    def _send_all_to_export(self):
        all_rows = self.db.get_all()
        if not all_rows:
            QMessageBox.information(self, "Empty", "No records in the database yet.")
            return
        if self._export_tab:
            self._export_tab.load_subjects(all_rows)
            QMessageBox.information(self, "Entire Database Sent",
                                    f"All {len(all_rows)} patient records sent to Export.")

#  TAB 4 — EXPORT

class ExportWorker(QThread):
    progress = Signal(str, int)
    finished = Signal(int, list)

    def __init__(self, subjects, dest, exts, do_excel):
        super().__init__()
        self.subjects = subjects
        self.dest = dest
        self.exts = exts
        self.do_excel = do_excel

    def run(self):
        total = len(self.subjects)
        copied = 0
        failed = []
        for i, subj in enumerate(self.subjects):
            name = subj.get("folder_name", "")
            self.progress.emit(f"Copying {i+1}/{total}: {name}",
                               int((i / total) * 90))
            try:
                src = Path(subj.get("folder_path", ""))
                if not src.exists():
                    failed.append((name, "Source not found"))
                    continue
                out = self.dest / src.name
                out.mkdir(parents=True, exist_ok=True)
                for f in src.iterdir():
                    if f.is_file() and (self.exts is None or
                                        f.suffix.lower() in self.exts):
                        shutil.copy2(str(f), str(out / f.name))
                copied += 1
            except Exception as e:
                failed.append((name, str(e)))

        if self.do_excel:
            try:
                self.progress.emit("Creating Excel summary…", 95)
                create_export_excel(self.subjects, self.dest / "export_summary.xlsx")
            except Exception as e:
                failed.append(("Excel summary", str(e)))

        self.finished.emit(copied, failed)

class ExportTab(QWidget):
    def __init__(self, db: Database, parent=None):
        super().__init__(parent)
        self.db = db
        self._subjects: list[dict] = []
        self._dest: Path | None = None
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        layout.addWidget(make_page_header("Export",
                                          "Copy data from the original database for further analysis"))

        content = QWidget()
        clay = QVBoxLayout(content)
        clay.setContentsMargins(16, 12, 16, 12)
        clay.setSpacing(10)

        # Subjects group
        subj_group = QGroupBox("Subjects to Export")
        sglay = QVBoxLayout(subj_group)

        self.subj_count_lbl = QLabel("No subjects loaded — use the Search tab first.")
        sglay.addWidget(self.subj_count_lbl)

        self.subj_table = QTableWidget()
        self.subj_table.setColumnCount(3)
        self.subj_table.setHorizontalHeaderLabels(["Folder", "Name", "Exam Date"])
        self.subj_table.setAlternatingRowColors(True)
        self.subj_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.subj_table.verticalHeader().setVisible(False)
        self.subj_table.horizontalHeader().setStretchLastSection(True)
        self.subj_table.setMaximumHeight(200)
        sglay.addWidget(self.subj_table)

        rm_btn = QPushButton("Remove Selected")
        rm_btn.clicked.connect(self._remove_selected)
        sglay.addWidget(rm_btn, alignment=Qt.AlignRight)
        clay.addWidget(subj_group)

        # Options
        opts_group = QGroupBox("Export Options")
        oglay = QGridLayout(opts_group)

        oglay.addWidget(QLabel("Destination Folder:"), 0, 0)
        self.dest_lbl = QLabel("No folder selected")
        self.dest_lbl.setProperty("cssClass", "muted")
        self.dest_lbl.style().unpolish(self.dest_lbl)
        self.dest_lbl.style().polish(self.dest_lbl)
        self.dest_lbl.setWordWrap(True)
        oglay.addWidget(self.dest_lbl, 0, 1)
        browse_btn = QPushButton("Browse…")
        browse_btn.clicked.connect(self._browse_dest)
        oglay.addWidget(browse_btn, 0, 2)

        oglay.addWidget(QLabel("File Types:"), 1, 0)
        ft_row = QHBoxLayout()
        self.copy_c3d = QCheckBox("C3D (.c3d)")
        self.copy_c3d.setChecked(True)
        ft_row.addWidget(self.copy_c3d)
        self.copy_pdf = QCheckBox("PDF (.pdf)")
        self.copy_pdf.setChecked(True)
        ft_row.addWidget(self.copy_pdf)
        self.copy_all = QCheckBox("All files")
        ft_row.addWidget(self.copy_all)
        ft_row.addStretch()
        oglay.addLayout(ft_row, 1, 1, 1, 2)

        self.make_excel = QCheckBox("Create Excel summary (export_summary.xlsx)")
        self.make_excel.setChecked(True)
        oglay.addWidget(self.make_excel, 2, 0, 1, 3)

        clay.addWidget(opts_group)

        # Run
        self.export_btn = make_accent_btn("▶   Run Export")
        self.export_btn.clicked.connect(self._run_export)
        clay.addWidget(self.export_btn, alignment=Qt.AlignLeft)

        # Progress
        prog_group = QGroupBox("Progress")
        pglay = QVBoxLayout(prog_group)
        self.exp_status = QLabel("Ready.")
        pglay.addWidget(self.exp_status)
        self.exp_pbar = QProgressBar()
        self.exp_pbar.setMaximum(100)
        pglay.addWidget(self.exp_pbar)
        clay.addWidget(prog_group)

        clay.addStretch()
        layout.addWidget(content)

    def load_subjects(self, subjects: list[dict]):
        self._subjects = subjects
        self.subj_table.setRowCount(len(subjects))
        for i, r in enumerate(subjects):
            name = f"{r.get('last_name','')} {r.get('first_name','')}".strip()
            self.subj_table.setItem(i, 0, QTableWidgetItem(r.get("folder_name", "")))
            self.subj_table.setItem(i, 1, QTableWidgetItem(name))
            self.subj_table.setItem(i, 2, QTableWidgetItem(r.get("exam_date", "")))
        self.subj_count_lbl.setText(f"{len(subjects)} subject(s) loaded.")

    def _remove_selected(self):
        rows = sorted(set(idx.row() for idx in self.subj_table.selectedIndexes()), reverse=True)
        for row in rows:
            self.subj_table.removeRow(row)
            if row < len(self._subjects):
                self._subjects.pop(row)
        self.subj_count_lbl.setText(f"{len(self._subjects)} subject(s) loaded.")

    def _browse_dest(self):
        path = QFileDialog.getExistingDirectory(self, "Select Export Destination")
        if path:
            self._dest = Path(path)
            self.dest_lbl.setText(str(self._dest))

    def _run_export(self):
        if not self._subjects:
            QMessageBox.warning(self, "No Subjects", "No subjects loaded.")
            return
        if not self._dest:
            QMessageBox.critical(self, "No Destination", "Select a destination folder.")
            return
        exts = None
        if not self.copy_all.isChecked():
            exts = set()
            if self.copy_c3d.isChecked():
                exts.add(".c3d")
            if self.copy_pdf.isChecked():
                exts.add(".pdf")
            if not exts:
                QMessageBox.critical(self, "No File Types", "Select at least one file type.")
                return

        self.export_btn.setEnabled(False)
        self._export_worker = ExportWorker(
            self._subjects[:], self._dest, exts, self.make_excel.isChecked())
        self._export_worker.progress.connect(
            lambda msg, pct: (self.exp_status.setText(msg), self.exp_pbar.setValue(pct)))
        self._export_worker.finished.connect(self._export_done)
        self._export_worker.start()

    def _export_done(self, copied, failed):
        self.exp_pbar.setValue(100)
        self.exp_status.setText(f"Done — {copied} folders copied, {len(failed)} errors.")
        self.export_btn.setEnabled(True)
        if failed:
            QMessageBox.warning(self, "Export Complete",
                f"{copied} copied.\n{len(failed)} errors:\n" +
                "\n".join(f"  • {n}: {r}" for n, r in failed[:5]))
        else:
            QMessageBox.information(self, "Export Complete",
                f"✓  {copied} folders copied to:\n{self._dest}")

#  TAB 5 — ANONYMIZER

class _RedactionTermsManager:

    def __init__(self, parent: QWidget):
        self.parent = parent
        self.terms  = self._load()

    def _load(self) -> list[str]:
        try:
            if REDACTION_TERMS_FILE.exists():
                with open(REDACTION_TERMS_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
        except Exception:
            pass
        return REDACTION_TERMS_DEFAULT.copy()

    def save(self):
        try:
            with open(REDACTION_TERMS_FILE, "w", encoding="utf-8") as f:
                json.dump(self.terms, f, indent=2, ensure_ascii=False)
        except Exception as e:
            QMessageBox.critical(self.parent, "Save Error", str(e))

    def show_manager(self):
        self.dlg = QDialog(self.parent)
        self.dlg.setWindowTitle("Manage Redaction Terms")
        self.dlg.resize(600, 500)
        
        layout = QVBoxLayout(self.dlg)
        
        # Top row: Add + Search
        top_lay = QHBoxLayout()
        
        add_group = QGroupBox("Add New Term")
        alay = QHBoxLayout(add_group)
        self._new_term = QLineEdit()
        alay.addWidget(self._new_term)
        add_btn = QPushButton("Add")
        add_btn.clicked.connect(self._add)
        alay.addWidget(add_btn)
        top_lay.addWidget(add_group)
        
        srch_group = QGroupBox("Search")
        slay = QHBoxLayout(srch_group)
        self._srch_var = QLineEdit()
        slay.addWidget(self._srch_var)
        srch_btn = QPushButton("Find")
        srch_btn.clicked.connect(self._find)
        slay.addWidget(srch_btn)
        top_lay.addWidget(srch_group)
        
        layout.addLayout(top_lay)

        # Listbox
        layout.addWidget(QLabel("<b>Current Terms:</b>"))
        self.list_widget = QListWidget()
        self.list_widget.itemDoubleClicked.connect(self._edit)
        layout.addWidget(self.list_widget)

        # Actions
        act_lay = QHBoxLayout()
        edit_btn = QPushButton("Edit Selected")
        edit_btn.clicked.connect(self._edit)
        act_lay.addWidget(edit_btn)
        
        del_btn = QPushButton("Delete Selected")
        del_btn.clicked.connect(self._delete)
        act_lay.addWidget(del_btn)
        
        reset_btn = QPushButton("Reset to Defaults")
        reset_btn.clicked.connect(self._reset)
        act_lay.addWidget(reset_btn)
        act_lay.addStretch()
        layout.addLayout(act_lay)

        # Bottom Buttons
        bot_lay = QHBoxLayout()
        bot_lay.addStretch()
        save_btn = QPushButton("Save & Close")
        save_btn.setProperty("cssClass", "accent") # Assuming you use stylesheets
        save_btn.clicked.connect(self._save_and_close)
        bot_lay.addWidget(save_btn)
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.dlg.reject)
        bot_lay.addWidget(cancel_btn)
        
        layout.addLayout(bot_lay)

        self._refresh_list()
        self.dlg.exec()

    def _refresh_list(self):
        self.list_widget.clear()
        self.list_widget.addItems(sorted(self.terms))

    def _add(self):
        term = self._new_term.text().strip()
        if not term: return
        if term not in self.terms:
            self.terms.append(term)
        self._new_term.clear()
        self._refresh_list()

    def _find(self):
        q = self._srch_var.text().strip().lower()
        if not q: return
        items = self.list_widget.findItems(q, Qt.MatchFlag.MatchContains)
        if items:
            self.list_widget.setCurrentItem(items[0])
            self.list_widget.scrollToItem(items[0])

    def _edit(self):
        item = self.list_widget.currentItem()
        if not item: return
        old_term = item.text()
        
        # Native Qt Input Dialog makes this much easier than Tkinter
        new_term, ok = QInputDialog.getText(self.dlg, "Edit Term", "Edit term:", text=old_term)
        new_term = new_term.strip()
        
        if ok and new_term and new_term != old_term and new_term not in self.terms:
            self.terms[self.terms.index(old_term)] = new_term
            self._refresh_list()

    def _delete(self):
        item = self.list_widget.currentItem()
        if not item: return
        term = item.text()
        reply = QMessageBox.question(self.dlg, "Delete", f"Delete '{term}'?")
        if reply == QMessageBox.StandardButton.Yes:
            self.terms.remove(term)
            self._refresh_list()

    def _reset(self):
        reply = QMessageBox.question(self.dlg, "Reset", "Restore default terms? Custom changes will be lost.")
        if reply == QMessageBox.StandardButton.Yes:
            self.terms = REDACTION_TERMS_DEFAULT.copy()
            self._refresh_list()

    def _save_and_close(self):
        self.save()
        self.dlg.accept()

class AnonymizeWorker(QThread):
    progress = Signal(str, int)       # (status_text, percent)
    finished = Signal(bool, str)      # (success, message)

    def __init__(self, inp: str, out: str, terms: list[str], mode: str):
        super().__init__()
        self.inp = inp
        self.out = out
        self.terms = terms
        self.mode = mode

    def run(self):
        try:
            self.progress.emit("Opening PDF…", 10)
            with open(self.inp, "rb") as fh:
                data = fh.read()
            doc = fitz.open(stream=data, filetype="pdf")
            total = len(doc)

            all_terms = [t.strip() for t in self.terms if t.strip()]
            single_set = {t.lower() for t in all_terms if " " not in t}
            phrase_terms = [t for t in all_terms if " " in t]

            for pi, page in enumerate(doc):
                self.progress.emit(f"Page {pi+1}/{total}…",
                                   10 + int((pi / max(total, 1)) * 80))

                # Single-word redaction
                for (x0, y0, x1, y1, text, *_) in (page.get_text("words") or []):
                    if text.strip(",:;").lower() in single_set:
                        r = fitz.Rect(x0 - 0.5, y0 - 0.5, x1 + 0.5, y1 + 0.5)
                        page.add_redact_annot(r, text="[REDACTED]", fill=(0, 0, 0))
                        if self.mode == "aggressive":
                            page.add_redact_annot(
                                fitz.Rect(r.x1, r.y0-2, r.x1+150, r.y1+2),
                                text=" ", fill=(0, 0, 0))
                            page.add_redact_annot(
                                fitz.Rect(r.x0-10, r.y1, r.x1+150, r.y1+20),
                                text=" ", fill=(0, 0, 0))

                # Phrase redaction
                for phrase in phrase_terms:
                    for r in page.search_for(phrase, flags=1):
                        page.add_redact_annot(r, text="[REDACTED]", fill=(0, 0, 0))
                        if self.mode == "aggressive":
                            page.add_redact_annot(
                                fitz.Rect(r.x1, r.y0-2, r.x1+150, r.y1+2),
                                text=" ", fill=(0, 0, 0))

                page.apply_redactions()

            self.progress.emit("Saving…", 95)
            doc.save(self.out)
            doc.close()
            self.finished.emit(True, self.out)
        except Exception as e:
            self.finished.emit(False, str(e))

class AnonymizerTab(QWidget):

    def __init__(self, get_selected_cb, parent=None):
        super().__init__(parent)
        self.get_selected = get_selected_cb
        self.terms_mgr = _RedactionTermsManager(self)
        self._input_file = ""
        self._output_folder = str(Path.home())
        self._output_fname = "anonymized_document.pdf"
        self._redact_mode = "standard"
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        layout.addWidget(make_page_header(
            "Clinical Anonymizer",
            "Remove sensitive parts of clinical reports before interpretation"))

        content = QWidget()
        clay = QVBoxLayout(content)
        clay.setContentsMargins(16, 12, 16, 12)
        clay.setSpacing(10)

        # Patient integration bar
        pi_row = QHBoxLayout()
        lbl = QLabel("Selected patient:")
        lbl.setStyleSheet("font-weight: 700;")
        pi_row.addWidget(lbl)
        self.patient_lbl = QLabel("None — select a patient in the Patients tab")
        self.patient_lbl.setProperty("cssClass", "muted")
        self.patient_lbl.style().unpolish(self.patient_lbl)
        self.patient_lbl.style().polish(self.patient_lbl)
        pi_row.addWidget(self.patient_lbl, stretch=1)
        use_pdf_btn = QPushButton("📄  Use This Patient's PDF")
        use_pdf_btn.clicked.connect(self._use_patient_pdf)
        pi_row.addWidget(use_pdf_btn)
        clay.addLayout(pi_row)

        # 1. Document configuration
        io_group = QGroupBox("1.  Document Configuration")
        igrid = QGridLayout(io_group)

        igrid.addWidget(QLabel("Source PDF:"), 0, 0)
        self.input_edit = QLineEdit()
        self.input_edit.setReadOnly(True)
        self.input_edit.setPlaceholderText("No file selected…")
        igrid.addWidget(self.input_edit, 0, 1)
        brw_in = QPushButton("Browse…")
        brw_in.clicked.connect(self._browse_input)
        igrid.addWidget(brw_in, 0, 2)

        igrid.addWidget(QLabel("Destination:"), 1, 0)
        self.output_dir_edit = QLineEdit()
        self.output_dir_edit.setReadOnly(True)
        self.output_dir_edit.setText(str(Path.home()))
        igrid.addWidget(self.output_dir_edit, 1, 1)
        brw_out = QPushButton("Browse…")
        brw_out.clicked.connect(self._browse_output)
        igrid.addWidget(brw_out, 1, 2)

        igrid.addWidget(QLabel("Filename:"), 2, 0)
        self.fname_edit = QLineEdit("anonymized_document.pdf")
        igrid.addWidget(self.fname_edit, 2, 1, 1, 2)

        igrid.setColumnStretch(1, 1)
        clay.addWidget(io_group)

        # 2. Redaction configuration
        cfg_group = QGroupBox("2.  Redaction Configuration")
        cfg_lay = QHBoxLayout(cfg_group)

        # Left: methodology
        method_w = QWidget()
        mlay = QVBoxLayout(method_w)
        mlay.setContentsMargins(0, 0, 0, 0)
        ml = QLabel("Methodology:")
        ml.setStyleSheet("font-weight: 700;")
        mlay.addWidget(ml)
        self.radio_standard = QRadioButton(
            "Standard Protection (redacts identified words)")
        self.radio_standard.setChecked(True)
        mlay.addWidget(self.radio_standard)
        self.radio_aggressive = QRadioButton(
            "Enhanced Protection (redacts contextual area)")
        mlay.addWidget(self.radio_aggressive)
        self._mode_group = QButtonGroup(self)
        self._mode_group.addButton(self.radio_standard)
        self._mode_group.addButton(self.radio_aggressive)
        mlay.addStretch()
        cfg_lay.addWidget(method_w)

        # Right: terms
        terms_w = QWidget()
        tlay = QVBoxLayout(terms_w)
        tlay.setContentsMargins(0, 0, 0, 0)
        tl = QLabel("Sensitive Terms:")
        tl.setStyleSheet("font-weight: 700;")
        tlay.addWidget(tl)
        tsub = QLabel("Labels used to detect PHI in the document:")
        tsub.setProperty("cssClass", "muted")
        tsub.style().unpolish(tsub)
        tsub.style().polish(tsub)
        tlay.addWidget(tsub)
        mgr_btn = make_accent_btn("⚙  Manage Redaction Terms")
        mgr_btn.clicked.connect(self.terms_mgr.show_manager)
        tlay.addWidget(mgr_btn, alignment=Qt.AlignLeft)
        tlay.addStretch()
        cfg_lay.addWidget(terms_w)

        clay.addWidget(cfg_group)

        # 3. Execute
        exec_row = QHBoxLayout()
        self.anon_btn = make_accent_btn("▶  Anonymize PDF")
        self.anon_btn.clicked.connect(self._start)
        exec_row.addWidget(self.anon_btn)
        rst_btn = QPushButton("⟳  Reset")
        rst_btn.clicked.connect(self._reset)
        exec_row.addWidget(rst_btn)
        exec_row.addStretch()
        clay.addLayout(exec_row)

        if fitz is None:
            warn = QLabel("⚠  PyMuPDF not installed — pip install pymupdf")
            warn.setStyleSheet(f"color: {PALETTE['warning']}; font-weight: 700;")
            clay.addWidget(warn)

        # Progress
        prog_group = QGroupBox("Progress")
        pglay = QVBoxLayout(prog_group)
        self.status_lbl = QLabel("Ready to anonymise.")
        pglay.addWidget(self.status_lbl)
        self.pbar = QProgressBar()
        self.pbar.setMaximum(100)
        pglay.addWidget(self.pbar)
        clay.addWidget(prog_group)

        clay.addStretch()

        # Wrap in scroll area for small screens (14")
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setWidget(content)
        scroll.setStyleSheet("QScrollArea { border: none; }")
        layout.addWidget(scroll)

    # patient integration
    def update_patient_label(self, rec: dict | None):
        if rec:
            name = (f"{rec.get('last_name','')} {rec.get('first_name','')}".strip()
                    or rec.get("folder_name", ""))
            self.patient_lbl.setText(f"{name}  —  {rec.get('folder_name','')}")
        else:
            self.patient_lbl.setText("None — select a patient in the Patients tab")

    def _use_patient_pdf(self):
        rec = self.get_selected()
        if not rec:
            QMessageBox.information(self, "No Patient",
                                    "Select a patient in the Patients tab first.")
            return
        pdf = rec.get("pdf_path", "")
        if not pdf or not Path(pdf).exists():
            QMessageBox.warning(self, "No PDF",
                                "No qualifying PDF found for this patient.")
            return
        self.input_edit.setText(pdf)
        self.fname_edit.setText(f"{Path(pdf).stem}_anonymized.pdf")
        self.output_dir_edit.setText(str(Path(pdf).parent))

    # dialogs
    def _browse_input(self):
        f, _ = QFileDialog.getOpenFileName(
            self, "Select PDF",
            filter="PDF Files (*.pdf);;All Files (*.*)")
        if f:
            self.input_edit.setText(f)
            self.fname_edit.setText(f"{Path(f).stem}_anonymized.pdf")
            self.output_dir_edit.setText(str(Path(f).parent))

    def _browse_output(self):
        d = QFileDialog.getExistingDirectory(self, "Select Output Folder")
        if d:
            self.output_dir_edit.setText(d)

    def _reset(self):
        self.input_edit.clear()
        self.output_dir_edit.setText(str(Path.home()))
        self.fname_edit.setText("anonymized_document.pdf")
        self.radio_standard.setChecked(True)
        self.pbar.setValue(0)
        self.status_lbl.setText("Ready to anonymise.")

    # worker
    def _start(self):
        inp = self.input_edit.text().strip()
        if not inp or not Path(inp).exists():
            QMessageBox.critical(self, "Error", "Select a valid PDF first.")
            return
        if fitz is None:
            QMessageBox.critical(self, "Missing Library",
                                 "PyMuPDF required — pip install pymupdf")
            return

        out = str(Path(self.output_dir_edit.text()) / self.fname_edit.text())
        mode = "aggressive" if self.radio_aggressive.isChecked() else "standard"

        self.anon_btn.setEnabled(False)
        self.pbar.setValue(0)

        self._worker = AnonymizeWorker(inp, out, self.terms_mgr.terms[:], mode)
        self._worker.progress.connect(
            lambda msg, pct: (self.status_lbl.setText(msg), self.pbar.setValue(pct)))
        self._worker.finished.connect(self._on_finished)
        self._worker.start()

    def _on_finished(self, success: bool, message: str):
        self.anon_btn.setEnabled(True)
        if success:
            self.pbar.setValue(100)
            self.status_lbl.setText("Anonymization complete.")
            QMessageBox.information(self, "Done", f"✓ Saved to:\n{message}")
        else:
            self.status_lbl.setText(f"Error: {message}")
            QMessageBox.critical(self, "Error", message)
        QTimer.singleShot(3000, lambda: (
            self.pbar.setValue(0),
            self.status_lbl.setText("Ready to anonymise.")))

#  TAB 6 — AI INTERPRETER

def _ai_extract_file(path: Path) -> tuple[str, list[tuple[bytes, str]]]:
    """Return (text, [(raw_bytes, mime), …]) from a file."""
    text = ""
    imgs: list[tuple[bytes, str]] = []
    try:
        suf = path.suffix.lower()
        if suf == ".docx" and _DocxDocument:
            doc  = _DocxDocument(str(path))
            text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
            with zipfile.ZipFile(str(path)) as z:
                for name in z.namelist():
                    if name.startswith("word/media/"):
                        raw  = z.read(name)
                        mime, _ = mimetypes.guess_type(f"x{Path(name).suffix}")
                        imgs.append((raw, mime or "image/png"))
        elif suf == ".pdf":
            text = ""
            if fitz:
                try:
                    d = fitz.open(str(path))
                    text = "\n".join(p.get_text("text") for p in d)
                    d.close()
                except Exception:
                    pass
        elif suf in {".png", ".jpg", ".jpeg", ".bmp"}:
            raw  = path.read_bytes()
            mime, _ = mimetypes.guess_type(str(path))
            imgs.append((raw, mime or "image/png"))
        else:
            text = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        print(f"[AI extract] {path.name}: {e}")
    return text.strip(), imgs

class AIWorker(QThread):
    progress = Signal(str, int)       # (status_text, percent)
    finished = Signal(bool, str)      # (success, message_or_path)

    def __init__(self, file_list: list[dict], header: dict,
                 save_path: str, api_key: str, prompt: str = ""):
        super().__init__()
        self.file_list = file_list
        self.header = header
        self.save_path = save_path
        self.api_key = api_key
        self.prompt = prompt or CLINICAL_PROMPT

    def run(self):
        try:
            self.progress.emit("Connecting to OpenAI…", 10)
            client = _OpenAI(api_key=self.api_key)
            client.models.list()

            self.progress.emit("Extracting file contents…", 20)
            inputs: list[tuple[str, str, list]] = []
            for item in self.file_list:
                fp   = Path(item["file_path"])
                desc = item.get("description", "").strip()
                text, imgs = _ai_extract_file(fp)
                combined = f"Description: {desc}\n\n{text}".strip() if text \
                    else f"Description: {desc}"
                inputs.append((combined, fp.name, imgs))

            self.progress.emit("Building payload…", 35)
            content: list[dict] = [{"type": "text",
                                     "text": self.prompt.strip()}]
            for text, label, _ in inputs:
                content.append({"type": "text",
                                 "text": f"FILE [{label}]:\n{text}"})
            for _, _, imgs in inputs:
                for raw, mime in imgs:
                    content.append({
                        "type": "image_url",
                        "image_url": {
                            "url": (f"data:{mime};base64,"
                                    f"{base64.b64encode(raw).decode()}")
                        },
                    })

            self.progress.emit("Sending to AI (PHI excluded)…", 55)
            response = client.chat.completions.create(
                model="gpt-5",
                messages=[{"role": "user", "content": content}],
            )
            interpretation = response.choices[0].message.content.strip()
            self.progress.emit("Interpretation received.", 80)

            # Save report
            self.progress.emit("Saving report…", 90)
            header = self.header
            fname = (f"Report_{header.get('last_name','')}_"
                     f"{header.get('patient_id','')}_"
                     f"{datetime.now().strftime('%Y%m%d')}")
            if _DocxDocument:
                doc = _DocxDocument()
                doc.add_heading("Patient Report", level=0)
                doc.add_paragraph(f"Patient ID:  {header.get('patient_id','')}")
                doc.add_paragraph(f"Date:        {datetime.now().strftime('%Y-%m-%d')}")
                if header.get("diagnosis"):
                    doc.add_paragraph(f"Diagnosis context:  {header['diagnosis']}")
                doc.add_heading("AI Clinical Interpretation", level=1)
                for line in interpretation.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())
                final_path = str(Path(self.save_path) / (fname + ".docx"))
                doc.save(final_path)
            else:
                final_path = str(Path(self.save_path) / (fname + ".txt"))
                with open(final_path, "w", encoding="utf-8") as fh:
                    diag = header.get("diagnosis", "")
                    diag_line = f"Diagnosis context: {diag}\n\n" if diag else ""
                    fh.write(f"Patient ID: {header.get('patient_id','')}\n"
                             f"{diag_line}"
                             f"{interpretation}")

            self.finished.emit(True, final_path)
        except Exception as e:
            self.finished.emit(False, str(e))

class _APIKeyDialog(QDialog):

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("API Key & Consent")
        self.setMinimumSize(520, 460)
        self.resize(520, 460)
        self.setModal(True)
        self.api_key: str | None = None
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(10)

        title = QLabel("Legal & Privacy Warning")
        title.setStyleSheet(f"font-size: 14px; font-weight: 700; "
                            f"color: {PALETTE['primary']};")
        layout.addWidget(title)

        # Legal text
        legal_txt = QTextEdit()
        legal_txt.setReadOnly(True)
        legal_txt.setPlainText(LEGAL_WARNING)
        legal_txt.setStyleSheet("border: 1px solid " + PALETTE['border'] + ";"
                                " border-radius: 6px; padding: 8px;")
        legal_txt.setMaximumHeight(180)
        layout.addWidget(legal_txt)

        # API key
        key_row = QHBoxLayout()
        kl = QLabel("OpenAI API Key:")
        kl.setStyleSheet("font-weight: 700;")
        key_row.addWidget(kl)
        self.key_edit = QLineEdit()
        self.key_edit.setEchoMode(QLineEdit.Password)
        self.key_edit.setPlaceholderText("sk-...")
        key_row.addWidget(self.key_edit)
        layout.addLayout(key_row)

        # Help link
        help_btn = QPushButton("How to get an API key?")
        help_btn.setFlat(True)
        help_btn.setStyleSheet(f"color: {PALETTE['accent']}; text-align: left;")
        help_btn.setCursor(Qt.PointingHandCursor)
        help_btn.clicked.connect(self._show_help)
        layout.addWidget(help_btn, alignment=Qt.AlignLeft)

        layout.addStretch()

        # Buttons
        btn_row = QHBoxLayout()
        btn_row.addStretch()
        ok_btn = make_accent_btn("OK && Process")
        ok_btn.clicked.connect(self._ok)
        btn_row.addWidget(ok_btn)
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        btn_row.addWidget(cancel_btn)
        layout.addLayout(btn_row)

        self.key_edit.returnPressed.connect(self._ok)

    def _ok(self):
        key = self.key_edit.text().strip()
        if not key:
            QMessageBox.warning(self, "Missing Key", "Enter your API key.")
            return
        if not (key.startswith("sk-") and len(key) > 20):
            reply = QMessageBox.question(
                self, "Invalid?", "Key format looks odd. Proceed anyway?")
            if reply != QMessageBox.StandardButton.Yes:
                return
        self.api_key = key
        self.accept()

    def _show_help(self):
        dlg = QDialog(self)
        dlg.setWindowTitle("API Key Help")
        dlg.resize(440, 280)
        dlg.setModal(True)
        lay = QVBoxLayout(dlg)
        txt = QTextEdit()
        txt.setReadOnly(True)
        txt.setPlainText(API_HELP)
        lay.addWidget(txt)
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dlg.accept)
        lay.addWidget(close_btn, alignment=Qt.AlignRight)
        dlg.exec()

    def get_key(self) -> str | None:
        """Show the dialog modally and return the API key, or None if cancelled."""
        self.exec()
        return self.api_key

class _AddFileDialog(QDialog):

    def __init__(self, parent=None, initial_path: str = ""):
        super().__init__(parent)
        self.setWindowTitle("Add File")
        self.setMinimumWidth(520)
        self.resize(560, 200)
        self.setModal(True)
        self.result_path = ""
        self.result_desc = ""
        self._build(initial_path)

    def _build(self, initial_path: str):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(10)

        # File path
        fp_row = QHBoxLayout()
        fp_row.addWidget(QLabel("File:"))
        self._path_edit = QLineEdit(initial_path)
        self._path_edit.setReadOnly(True)
        self._path_edit.setPlaceholderText("No file selected…")
        fp_row.addWidget(self._path_edit, stretch=1)
        brw = QPushButton("Browse…")
        brw.clicked.connect(self._browse)
        fp_row.addWidget(brw)
        layout.addLayout(fp_row)

        # Auto-detected type
        self._type_lbl = QLabel("")
        self._type_lbl.setStyleSheet(f"color: {PALETTE['text_muted']}; font-size: 12px;")
        layout.addWidget(self._type_lbl)

        # Description
        layout.addWidget(QLabel("Description (sent to AI as context):"))
        self._desc_edit = QLineEdit()
        self._desc_edit.setPlaceholderText(
            "e.g., Anonymized gait report, Feature extraction output, "
            "Pre-surgery kinematics…")
        layout.addWidget(self._desc_edit)

        layout.addStretch()

        # Buttons
        btn_row = QHBoxLayout()
        btn_row.addStretch()
        ok_btn = make_accent_btn("Add")
        ok_btn.clicked.connect(self._ok)
        btn_row.addWidget(ok_btn)
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        btn_row.addWidget(cancel_btn)
        layout.addLayout(btn_row)

        if initial_path:
            self._update_type(initial_path)

    def _browse(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Select file",
            filter="All supported (*.pdf *.xlsx *.txt *.docx *.png *.jpg *.jpeg);;"
                   "PDF (*.pdf);;Excel (*.xlsx);;Text (*.txt);;"
                   "Word (*.docx);;Images (*.png *.jpg *.jpeg);;All (*.*)")
        if path:
            self._path_edit.setText(path)
            self._update_type(path)
            if not self._desc_edit.text():
                self._desc_edit.setText(Path(path).stem)

    def _update_type(self, path: str):
        ext = Path(path).suffix.lower()
        type_map = {
            ".pdf": "📄 PDF document",
            ".xlsx": "📊 Excel spreadsheet (features / C3D data)",
            ".txt": "📝 Text file (LLM feature summary)",
            ".docx": "📃 Word document",
            ".png": "🖼 Image (PNG)", ".jpg": "🖼 Image (JPEG)",
            ".jpeg": "🖼 Image (JPEG)",
        }
        self._type_lbl.setText(f"Type: {type_map.get(ext, f'Unknown ({ext})')}")

    def _ok(self):
        if not self._path_edit.text():
            QMessageBox.warning(self, "No File", "Select a file first.")
            return
        self.result_path = self._path_edit.text()
        self.result_desc = self._desc_edit.text().strip() or Path(self.result_path).stem
        self.accept()

class _PromptEditorDialog(QDialog):

    def __init__(self, parent=None, current_text: str = "",
                 current_name: str = ""):
        super().__init__(parent)
        self.setWindowTitle("Prompt Editor")
        self.setMinimumSize(680, 520)
        self.resize(780, 620)
        self.setModal(True)
        self.result_text: str | None = None
        self._current_name = current_name
        self._build(current_text)

    def _build(self, text: str):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        # Header
        hdr = QFrame()
        hdr.setFixedHeight(48)
        hdr.setStyleSheet(f"background-color: {PALETTE['primary']};")
        hlay = QHBoxLayout(hdr)
        hlay.setContentsMargins(20, 0, 20, 0)
        t = QLabel("✎  Prompt Editor")
        t.setStyleSheet("color: white; font-size: 14px; font-weight: 700;")
        hlay.addWidget(t)
        layout.addWidget(hdr)

        body = QVBoxLayout()
        body.setContentsMargins(16, 12, 16, 12)
        body.setSpacing(8)

        # Preset loading
        preset_row = QHBoxLayout()
        preset_row.addWidget(QLabel("Load preset:"))
        for key, data in PROMPT_PRESETS.items():
            btn = QPushButton(f"{data['icon']}  {key}")
            btn.setToolTip(data["description"])
            btn.clicked.connect(lambda checked, k=key: self._load_preset(k))
            preset_row.addWidget(btn)
        preset_row.addStretch()
        body.addLayout(preset_row)

        # Custom prompt loading
        customs = _load_custom_prompts()
        if customs:
            custom_row = QHBoxLayout()
            custom_row.addWidget(QLabel("Custom:"))
            for name in customs:
                btn = QPushButton(f"📝 {name}")
                btn.clicked.connect(
                    lambda checked, n=name: self._load_custom(n))
                custom_row.addWidget(btn)
            custom_row.addStretch()
            body.addLayout(custom_row)

        # Editor
        self._editor = QPlainTextEdit()
        self._editor.setPlainText(text)
        self._editor.setStyleSheet(
            f"font-family: 'Consolas', 'Fira Code', monospace; font-size: 12px;"
            f" background-color: {PALETTE['surface']};"
            f" border: 1px solid {PALETTE['border']}; border-radius: 6px;"
            f" padding: 8px;")
        body.addWidget(self._editor, stretch=1)

        # Info
        info = QLabel("The prompt is sent as the system instruction to the AI. "
                       "The patient data files are appended automatically.")
        info.setWordWrap(True)
        info.setStyleSheet(f"color: {PALETTE['text_muted']}; font-size: 11px;")
        body.addWidget(info)

        layout.addLayout(body)

        # Button bar
        btn_bar = QWidget()
        btn_bar.setStyleSheet(
            f"background-color: {PALETTE['bg']};"
            f" border-top: 1px solid {PALETTE['border']};")
        blay = QHBoxLayout(btn_bar)
        blay.setContentsMargins(16, 10, 16, 10)

        save_btn = QPushButton("💾  Save as Custom…")
        save_btn.clicked.connect(self._save_custom)
        blay.addWidget(save_btn)

        del_btn = QPushButton("🗑  Delete Custom…")
        del_btn.clicked.connect(self._delete_custom)
        blay.addWidget(del_btn)

        blay.addStretch()

        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        blay.addWidget(cancel_btn)

        ok_btn = make_accent_btn("Apply Prompt")
        ok_btn.clicked.connect(self._apply)
        blay.addWidget(ok_btn)

        layout.addWidget(btn_bar)

    def _load_preset(self, key: str):
        self._editor.setPlainText(PROMPT_PRESETS[key]["prompt"])
        self._current_name = key

    def _load_custom(self, name: str):
        customs = _load_custom_prompts()
        if name in customs:
            self._editor.setPlainText(customs[name]["prompt"])
            self._current_name = name

    def _save_custom(self):
        name, ok = QInputDialog.getText(
            self, "Save Custom Prompt",
            "Enter a name for this prompt:",
            text=self._current_name or "My Custom Prompt")
        if ok and name.strip():
            name = name.strip()
            if name in PROMPT_PRESETS:
                QMessageBox.warning(self, "Reserved Name",
                                    f'"{name}" is a built-in preset. Choose a different name.')
                return
            _save_custom_prompt(name, self._editor.toPlainText())
            self._current_name = name
            QMessageBox.information(self, "Saved",
                                    f'Prompt "{name}" saved. It will appear in the editor next time.')

    def _delete_custom(self):
        customs = _load_custom_prompts()
        if not customs:
            QMessageBox.information(self, "No Custom Prompts",
                                    "No custom prompts saved yet.")
            return
        name, ok = QInputDialog.getItem(
            self, "Delete Custom Prompt",
            "Select prompt to delete:", list(customs.keys()), 0, False)
        if ok and name:
            _delete_custom_prompt(name)
            QMessageBox.information(self, "Deleted", f'Prompt "{name}" deleted.')

    def _apply(self):
        self.result_text = self._editor.toPlainText()
        self.accept()

class AIInterpreterTab(QWidget):

    def __init__(self, get_selected_cb, parent=None):
        super().__init__(parent)
        self.get_selected = get_selected_cb
        self.file_list: list[dict] = []  # [{file_path, description}, …]
        self.save_location = ""
        self._active_prompt_name = DEFAULT_PROMPT_KEY
        self._active_prompt_text = PROMPT_PRESETS[DEFAULT_PROMPT_KEY]["prompt"]
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        layout.addWidget(make_page_header(
            "AI Clinical Interpreter",
            "Send anonymized files to AI for structured clinical interpretation"))

        content = QWidget()
        clay = QVBoxLayout(content)
        clay.setContentsMargins(16, 12, 16, 12)
        clay.setSpacing(10)

        # Patient integration bar
        pi_row = QHBoxLayout()
        pi_row.addWidget(QLabel("Selected patient:"))
        self.patient_lbl = QLabel("None — select a patient in the Patients tab")
        self.patient_lbl.setProperty("cssClass", "muted")
        self.patient_lbl.style().unpolish(self.patient_lbl)
        self.patient_lbl.style().polish(self.patient_lbl)
        pi_row.addWidget(self.patient_lbl, stretch=1)
        load_btn = QPushButton("📂  Load from Patient Folder")
        load_btn.clicked.connect(self._load_from_patient)
        pi_row.addWidget(load_btn)
        clay.addLayout(pi_row)

        # 1. Input Files
        file_group = QGroupBox(
            "1.  Input Files (anonymized PDFs, feature outputs, images, etc.)")
        fglay = QVBoxLayout(file_group)

        fb_row = QHBoxLayout()
        add_btn = make_accent_btn("+ Add File…")
        add_btn.clicked.connect(self._add_file)
        fb_row.addWidget(add_btn)
        rm_btn = QPushButton("Remove Selected")
        rm_btn.clicked.connect(self._remove_selected)
        fb_row.addWidget(rm_btn)
        clr_btn = QPushButton("Clear All")
        clr_btn.clicked.connect(self._clear_files)
        fb_row.addWidget(clr_btn)
        fb_row.addStretch()
        fglay.addLayout(fb_row)

        self.file_table = QTableWidget()
        self.file_table.setColumnCount(4)
        self.file_table.setHorizontalHeaderLabels(
            ["Filename", "Type", "Description", "Size"])
        self.file_table.setAlternatingRowColors(True)
        self.file_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.file_table.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.file_table.verticalHeader().setVisible(False)
        self.file_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.file_table.horizontalHeader().setStretchLastSection(True)
        self.file_table.setColumnWidth(0, 180)
        self.file_table.setColumnWidth(1, 70)
        self.file_table.setColumnWidth(2, 250)
        self.file_table.setMaximumHeight(140)
        fglay.addWidget(self.file_table)

        self._file_count_lbl = QLabel("No files added.")
        self._file_count_lbl.setStyleSheet(
            f"color: {PALETTE['text_muted']}; font-size: 12px;")
        fglay.addWidget(self._file_count_lbl)
        clay.addWidget(file_group)

        # 2. Prompt Configuration
        prompt_group = QGroupBox("2.  Prompt Configuration")
        pglay = QVBoxLayout(prompt_group)

        # Preset buttons
        preset_row = QHBoxLayout()
        preset_row.addWidget(QLabel("Preset:"))
        self._preset_btns: dict[str, QPushButton] = {}
        for key, data in PROMPT_PRESETS.items():
            btn = QPushButton(f"{data['icon']} {key}")
            btn.setToolTip(data["description"])
            btn.setCheckable(True)
            btn.clicked.connect(
                lambda checked, k=key: self._select_preset(k))
            self._preset_btns[key] = btn
            preset_row.addWidget(btn)
        preset_row.addStretch()
        pglay.addLayout(preset_row)
        self._preset_btns[DEFAULT_PROMPT_KEY].setChecked(True)

        # Diagnosis context — anchors the prompt in a specific disease
        diag_row = QVBoxLayout()
        diag_lbl_row = QHBoxLayout()
        diag_lbl = QLabel("Diagnosis context:")
        diag_lbl.setStyleSheet("font-weight: 700;")
        diag_lbl_row.addWidget(diag_lbl)
        diag_hint = QLabel(
            "(disease / condition the analysis should be framed in — "
            "e.g. 'spastic diplegia CP, GMFCS II', 'post-stroke left hemiparesis', "
            "'unilateral knee OA')")
        diag_hint.setStyleSheet(
            f"color: {PALETTE['text_muted']}; font-size: 11px;")
        diag_hint.setWordWrap(True)
        diag_lbl_row.addWidget(diag_hint, stretch=1)
        diag_row.addLayout(diag_lbl_row)
        self.diagnosis_edit = QPlainTextEdit()
        self.diagnosis_edit.setPlaceholderText(
            "Required — describe the diagnosis (one or two lines). "
            "All reasoning, cause-and-effect chains, and patient-vs-healthy "
            "comparisons will be framed within this context.")
        self.diagnosis_edit.setMaximumHeight(70)
        diag_row.addWidget(self.diagnosis_edit)
        pglay.addLayout(diag_row)

        # Edit / info row
        edit_row = QHBoxLayout()
        edit_btn = QPushButton("✎  Edit Prompt…")
        edit_btn.clicked.connect(self._edit_prompt)
        edit_row.addWidget(edit_btn)
        self._prompt_info_lbl = QLabel(
            f"Using: {DEFAULT_PROMPT_KEY}")
        self._prompt_info_lbl.setStyleSheet(
            f"color: {PALETTE['text_muted']}; font-size: 12px;")
        edit_row.addWidget(self._prompt_info_lbl, stretch=1)
        pglay.addLayout(edit_row)

        clay.addWidget(prompt_group)

        # 3. Report Output
        out_group = QGroupBox("3.  Report Output")
        oglay = QVBoxLayout(out_group)

        # Patient info row
        info_grid = QGridLayout()
        self.fn_edit = QLineEdit()
        self.fn_edit.setPlaceholderText("First name")
        self.ln_edit = QLineEdit()
        self.ln_edit.setPlaceholderText("Last name")
        self.id_edit = QLineEdit()
        self.id_edit.setPlaceholderText("Patient ID")
        self.dob_edit = QLineEdit()
        self.dob_edit.setPlaceholderText("Birthdate")
        for ci, (lbl, w) in enumerate([
            ("First:", self.fn_edit), ("Last:", self.ln_edit),
            ("ID:", self.id_edit), ("DOB:", self.dob_edit),
        ]):
            info_grid.addWidget(QLabel(lbl), 0, ci * 2)
            info_grid.addWidget(w, 0, ci * 2 + 1)
        for c in (1, 3, 5, 7):
            info_grid.setColumnStretch(c, 1)
        oglay.addLayout(info_grid)

        # Save location
        save_row = QHBoxLayout()
        save_row.addWidget(QLabel("Save to:"))
        self.save_lbl = QLabel("No location selected")
        self.save_lbl.setProperty("cssClass", "muted")
        self.save_lbl.style().unpolish(self.save_lbl)
        self.save_lbl.style().polish(self.save_lbl)
        save_row.addWidget(self.save_lbl, stretch=1)
        sv_btn = QPushButton("Browse…")
        sv_btn.clicked.connect(self._browse_save)
        save_row.addWidget(sv_btn)
        oglay.addLayout(save_row)

        clay.addWidget(out_group)

        # 4. Execute
        exec_row = QHBoxLayout()
        self.process_btn = make_accent_btn("▶  Send to AI & Generate Report")
        self.process_btn.clicked.connect(self._process)
        exec_row.addWidget(self.process_btn)
        reset_btn = QPushButton("⟳  Clear Form")
        reset_btn.clicked.connect(self._reset_form)
        exec_row.addWidget(reset_btn)
        exec_row.addStretch()
        clay.addLayout(exec_row)

        if _OpenAI is None:
            warn = QLabel("⚠  openai library not installed — pip install openai")
            warn.setStyleSheet(f"color: {PALETTE['warning']}; font-weight: 700;")
            clay.addWidget(warn)

        # Status
        stat_group = QGroupBox("Status")
        sglay = QVBoxLayout(stat_group)
        self.status_lbl = QLabel("Ready.")
        sglay.addWidget(self.status_lbl)
        self.pbar = QProgressBar()
        self.pbar.setMaximum(100)
        sglay.addWidget(self.pbar)
        clay.addWidget(stat_group)

        clay.addStretch()

        # Wrap in scroll area for small screens
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setWidget(content)
        scroll.setStyleSheet("QScrollArea { border: none; }")
        layout.addWidget(scroll)

    # Patient integration

    def update_patient_label(self, rec: dict | None):
        if rec:
            name = (f"{rec.get('last_name','')} {rec.get('first_name','')}".strip()
                    or rec.get("folder_name", ""))
            self.patient_lbl.setText(f"{name}  —  {rec.get('folder_name','')}")
            self.fn_edit.setText(rec.get("first_name", "") or "")
            self.ln_edit.setText(rec.get("last_name", "") or "")
            self.id_edit.setText(
                rec.get("ganglabor_id", "") or rec.get("folder_name", ""))
            self.dob_edit.setText(rec.get("birth_date", "") or "")
            # Auto-fill diagnosis context if the field is empty and the
            # patient record carries a diagnosis. The user can still edit it.
            patient_diag = (rec.get("diagnosis") or "").strip()
            if patient_diag and not self.diagnosis_edit.toPlainText().strip():
                self.diagnosis_edit.setPlainText(patient_diag)
        else:
            self.patient_lbl.setText(
                "None — select a patient in the Patients tab")

    def _load_from_patient(self):
        rec = self.get_selected()
        if not rec:
            QMessageBox.information(
                self, "No Patient",
                "Select a patient in the Patients tab first.")
            return
        folder = Path(rec.get("folder_path", ""))
        if not folder.exists():
            QMessageBox.warning(self, "No Folder", "Patient folder not found.")
            return
        added = 0
        for f in sorted(folder.iterdir()):
            if not f.is_file() or f.suffix.lower() in {".c3d", ".db"}:
                continue
            if any(item["file_path"] == str(f) for item in self.file_list):
                continue
            self.file_list.append({
                "file_path": str(f), "description": f.stem})
            added += 1
        self._refresh_file_table()
        QMessageBox.information(self, "Loaded", f"{added} file(s) loaded.")

    # File management

    def _add_file(self):
        dlg = _AddFileDialog(self)
        if dlg.exec():
            if any(item["file_path"] == dlg.result_path
                   for item in self.file_list):
                QMessageBox.information(
                    self, "Duplicate", "This file is already in the list.")
                return
            self.file_list.append({
                "file_path": dlg.result_path,
                "description": dlg.result_desc,
            })
            self._refresh_file_table()

    def _remove_selected(self):
        rows = sorted(
            set(idx.row() for idx in self.file_table.selectedIndexes()),
            reverse=True)
        for row in rows:
            if row < len(self.file_list):
                self.file_list.pop(row)
        self._refresh_file_table()

    def _clear_files(self):
        self.file_list.clear()
        self._refresh_file_table()

    def _refresh_file_table(self):
        self.file_table.setRowCount(len(self.file_list))
        ext_icons = {
            ".pdf": "PDF", ".xlsx": "XLSX", ".txt": "TXT",
            ".docx": "DOCX", ".png": "IMG", ".jpg": "IMG", ".jpeg": "IMG",
        }
        for i, item in enumerate(self.file_list):
            p = Path(item["file_path"])
            ext = p.suffix.lower()
            self.file_table.setItem(i, 0, QTableWidgetItem(p.name))
            self.file_table.setItem(
                i, 1, QTableWidgetItem(ext_icons.get(ext, ext)))
            self.file_table.setItem(
                i, 2, QTableWidgetItem(item.get("description", "")))
            try:
                size = p.stat().st_size
                if size > 1_000_000:
                    sz_str = f"{size / 1_000_000:.1f} MB"
                else:
                    sz_str = f"{size / 1000:.0f} KB"
            except Exception:
                sz_str = "?"
            self.file_table.setItem(i, 3, QTableWidgetItem(sz_str))

        n = len(self.file_list)
        self._file_count_lbl.setText(
            f"{n} file{'s' if n != 1 else ''} ready."
            if n else "No files added.")

    # Prompt management

    def _select_preset(self, key: str):
        self._active_prompt_name = key
        self._active_prompt_text = PROMPT_PRESETS[key]["prompt"]
        self._prompt_info_lbl.setText(f"Using: {key}")
        for k, btn in self._preset_btns.items():
            btn.setChecked(k == key)

    def _edit_prompt(self):
        dlg = _PromptEditorDialog(
            self, self._active_prompt_text, self._active_prompt_name)
        if dlg.exec() and dlg.result_text is not None:
            self._active_prompt_text = dlg.result_text
            self._active_prompt_name = "(edited)"
            self._prompt_info_lbl.setText("Using: Custom (edited)")
            for btn in self._preset_btns.values():
                btn.setChecked(False)

    # Save / Process

    def _browse_save(self):
        path = QFileDialog.getExistingDirectory(self, "Select Save Location")
        if path:
            self.save_location = path
            self.save_lbl.setText(f"Saving to: {path}")

    def _reset_form(self):
        self.fn_edit.clear()
        self.ln_edit.clear()
        self.id_edit.clear()
        self.dob_edit.clear()
        self.diagnosis_edit.clear()
        self.save_location = ""
        self.save_lbl.setText("No location selected")
        self.file_list.clear()
        self._refresh_file_table()
        self._select_preset(DEFAULT_PROMPT_KEY)
        self.status_lbl.setText("Ready.")
        self.pbar.setValue(0)

    def _process(self):
        if _OpenAI is None:
            QMessageBox.critical(
                self, "Missing Library",
                "openai not installed — pip install openai")
            return
        if not self.file_list:
            QMessageBox.warning(self, "No Files", "Add at least one file.")
            return
        if not self.save_location:
            QMessageBox.warning(
                self, "No Save Location", "Select a save location.")
            return

        # Diagnosis context — the prompt is anchored on this. Strongly
        # encouraged but not strictly required; warn the user if missing.
        diagnosis_text = self.diagnosis_edit.toPlainText().strip()
        if not diagnosis_text:
            reply = QMessageBox.question(
                self, "No Diagnosis Context",
                "You haven't entered a diagnosis. The reasoning will be "
                "much weaker without it — the AI will not know which "
                "disease to frame the cause-and-effect analysis around.\n\n"
                "Proceed anyway?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No)
            if reply != QMessageBox.StandardButton.Yes:
                return
            diagnosis_text = (
                "(diagnosis not specified — interpret cautiously, flag "
                "ambiguous findings, and avoid speculation about the "
                "underlying disease)")

        # Inject diagnosis into the {diagnosis_context} placeholder.
        # Using .replace() (not .format()) so other braces in the prompt
        # are left untouched.
        final_prompt = self._active_prompt_text.replace(
            "{diagnosis_context}", diagnosis_text)

        dlg = _APIKeyDialog(self)
        api_key = dlg.get_key()
        if not api_key:
            return

        header = {
            "first_name": self.fn_edit.text().strip(),
            "last_name":  self.ln_edit.text().strip(),
            "patient_id": self.id_edit.text().strip(),
            "birthdate":  self.dob_edit.text().strip(),
            "diagnosis":  diagnosis_text,
        }

        self.status_lbl.setText("Starting…")
        self.pbar.setValue(10)
        self.process_btn.setEnabled(False)

        self._ai_worker = AIWorker(
            self.file_list[:], header, self.save_location,
            api_key, final_prompt)
        self._ai_worker.progress.connect(
            lambda msg, pct: (
                self.status_lbl.setText(msg), self.pbar.setValue(pct)))
        self._ai_worker.finished.connect(self._on_ai_finished)
        self._ai_worker.start()

    def _on_ai_finished(self, success: bool, message: str):
        self.process_btn.setEnabled(True)
        if success:
            self.pbar.setValue(100)
            self.status_lbl.setText("Report saved.")
            QMessageBox.information(
                self, "Done", f"✓ Report saved to:\n{message}")
        else:
            self.status_lbl.setText(f"Error: {message}")
            QMessageBox.critical(self, "AI Error", message)
        QTimer.singleShot(3000, lambda: (
            self.status_lbl.setText("Ready."), self.pbar.setValue(0)))
